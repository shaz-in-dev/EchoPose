"""
EchoPose A-to-Z Visualized Guide Generator
Run: python scripts/generate_guide.py
Output: EchoPose_Guide.docx (project root)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)

# ── Colour palette ────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
C_BLUE   = RGBColor(0x1B, 0x6C, 0xA8)
C_TEAL   = RGBColor(0x00, 0x87, 0x8A)
C_GREEN  = RGBColor(0x1A, 0x73, 0x48)
C_RED    = RGBColor(0xC0, 0x20, 0x20)
C_ORANGE = RGBColor(0xD9, 0x6B, 0x00)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_LGREY  = RGBColor(0xF2, 0xF4, 0xF7)
C_MGREY  = RGBColor(0xCC, 0xD1, 0xDA)
C_DKGREY = RGBColor(0x3A, 0x3F, 0x4A)

# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    str(val.get('sz', 6)))
            el.set(qn('w:space'),'0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=C_NAVY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def add_para(text='', bold=False, italic=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def add_coloured_para(text, bg: RGBColor, fg: RGBColor = C_WHITE, size=11):
    """Single paragraph with coloured background (via table trick)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6.2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = True
    run.font.color.rgb = fg
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(8)
    doc.add_paragraph()  # spacing

def add_step_box(number: str, title: str, body_lines: list, color=C_BLUE):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = False
    # Number cell
    num_cell = tbl.cell(0, 0)
    num_cell.width = Cm(1.5)
    set_cell_bg(num_cell, color)
    np_ = num_cell.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    nr = np_.add_run(number)
    nr.font.size  = Pt(20)
    nr.font.bold  = True
    nr.font.color.rgb = C_WHITE
    # Body cell
    body_cell = tbl.cell(0, 1)
    set_cell_bg(body_cell, C_LGREY)
    bp = body_cell.paragraphs[0]
    br = bp.add_run(title)
    br.font.size  = Pt(12)
    br.font.bold  = True
    br.font.color.rgb = color
    for line in body_lines:
        lp = body_cell.add_paragraph(line)
        lp.paragraph_format.left_indent = Pt(4)
        for run in lp.runs:
            run.font.size = Pt(10)
    doc.add_paragraph()

def add_hardware_table(rows_data, headers):
    tbl = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = C_WHITE
        run.font.size  = Pt(10)
    # Data rows
    for ri, row_data in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        bg = C_LGREY if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            # Highlight cost column green
            if ci == len(row_data) - 1 and '£' in str(val):
                run.font.color.rgb = C_GREEN
                run.font.bold = True
    doc.add_paragraph()

def add_ascii_box(title: str, lines: list, border_color=C_TEAL):
    add_coloured_para(f'  DIAGRAM — {title}', border_color)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0x1A, 0x1A, 0x2E))
    bc = rgb_hex(border_color)
    set_cell_border(cell,
        top    = {'val': 'single', 'sz': 8, 'color': bc},
        bottom = {'val': 'single', 'sz': 8, 'color': bc},
        left   = {'val': 'single', 'sz': 8, 'color': bc},
        right  = {'val': 'single', 'sz': 8, 'color': bc},
    )
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0x00, 0xFF, 0xCC)
    doc.add_paragraph()

def add_warning(text, warn_type='NOTE'):
    colors = {
        'NOTE':    (C_BLUE,   '  ℹ  NOTE:    '),
        'WARNING': (C_ORANGE, '  ⚠  WARNING: '),
        'CRITICAL':(C_RED,    '  ✖  CRITICAL: '),
        'TIP':     (C_GREEN,  '  ✔  TIP:     '),
    }
    col, prefix = colors.get(warn_type, colors['NOTE'])
    add_coloured_para(prefix + text, col, C_WHITE, size=10)

def add_checklist(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'☐  {item}')
        run.font.size = Pt(10.5)

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('─' * 90)
    run.font.size  = Pt(7)
    run.font.color.rgb = C_MGREY

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

tbl = doc.add_table(rows=1, cols=1)
cell = tbl.cell(0, 0)
set_cell_bg(cell, C_NAVY)
cell.width = Inches(6.3)

for text, sz, bold, color in [
    ('ECHOPOSE', 42, True,  C_TEAL),
    ('WiFi CSI Pose Estimation System', 18, False, C_WHITE),
    ('', 8, False, C_WHITE),
    ('A-to-Z Setup, Training & Deployment Guide', 14, True, RGBColor(0xB0, 0xD0, 0xFF)),
    ('', 8, False, C_WHITE),
    ('From Hardware Purchase  →  Data Collection  →  Model Training', 11, False, C_MGREY),
    ('→  Product Build  →  Customer Deployment', 11, False, C_MGREY),
    ('', 8, False, C_WHITE),
    ('CONFIDENTIAL — COMMERCIAL IN CONFIDENCE', 9, True, C_ORANGE),
]:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size  = Pt(sz)
    run.font.bold  = bold
    run.font.color.rgb = color

doc.add_paragraph()
page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ══════════════════════════════════════════════════════════════════════════════

add_heading('TABLE OF CONTENTS', 1, C_NAVY)
toc_items = [
    ('PART 0', 'Hardware — What To Buy & Why',         3),
    ('PART 1', 'Physical Setup — Nodes, Router & Camera', 5),
    ('PART 2', 'Software Setup — PC Stack Installation', 8),
    ('PART 3', 'Data Collection — Recording Ground Truth', 10),
    ('PART 4', 'Model Training — Running & Evaluating',  13),
    ('PART 5', 'Product Build — Alerts, Dashboard & Installer', 16),
    ('PART 6', 'Deployment — SaaS & On-Premise',        19),
    ('APPENDIX A', 'Full Wiring Diagram',               22),
    ('APPENDIX B', 'Environment Variable Reference',    23),
    ('APPENDIX C', 'Troubleshooting Cheat Sheet',       24),
]
for part, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run1 = p.add_run(f'{part:<12}')
    run1.font.bold  = True
    run1.font.color.rgb = C_BLUE
    run1.font.size  = Pt(11)
    run2 = p.add_run(title)
    run2.font.size  = Pt(11)
    run3 = p.add_run(f'  ....  p{pg}')
    run3.font.color.rgb = C_MGREY
    run3.font.size  = Pt(10)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 0 — HARDWARE
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 0   HARDWARE — WHAT TO BUY & WHY', C_NAVY, C_WHITE, 14)
doc.add_paragraph()

add_heading('0.1  Depth Camera — Kinect Decision', 2, C_BLUE)

add_para('You asked: Xbox One Kinect 2  vs  Xbox 360 Kinect (from CeX)?', bold=True)
doc.add_paragraph()

add_hardware_table([
    ['Xbox One Kinect 2\n(RECOMMENDED ✔)', 'Time-of-Flight\n(true 3D depth)', '512×424 @ 30fps\ndepth + 1080p colour', 'USB 3.0\n+ Kinect Adapter', 'Best 3D skeleton\nground truth', '~£25–40'],
    ['Xbox 360 Kinect v1\n(NOT recommended)', 'Structured Light\n(older, noisier)', '640×480 @ 30fps\ndepth only', 'USB 2.0\n+ separate adapter', 'Less accurate\n2D-only reliable', '~£10–20'],
], ['Device', 'Depth Tech', 'Resolution / FPS', 'PC Connection', 'Quality', 'CeX Price'])

add_warning('The Xbox One Kinect 2 REQUIRES a "Kinect Adapter for Windows" to connect to a PC. Without it, it only works with an Xbox One. Search CeX for "Kinect Adapter" (Part: 6NU-00003). Costs ~£15–25. Buy both the sensor AND the adapter.', 'CRITICAL')

add_warning('Your Brio 500 webcam can be used as a 2D backup camera with Google MediaPipe (free). Useful for cross-checking. You already own this — use it from Day 1 while waiting for Kinect.', 'TIP')

divider()
add_heading('0.2  Router / Access Point Decision', 2, C_BLUE)
add_para('You are in student accommodation — you cannot replace the main router. You need your OWN isolated WiFi bubble that your ESP32 nodes connect to.', italic=True)
doc.add_paragraph()

add_hardware_table([
    ['TP-Link TL-WA801N\n(RECOMMENDED ✔)', 'Access Point mode\n(plug into room ethernet)', '2.4GHz only\n802.11n N300', '~£20', 'Simple, cheap, 2.4GHz\nExactly what ESP32 needs\nPlug-and-play AP mode'],
    ['Tenda AX1500 RX1 Pro', 'Router + AP mode', '2.4GHz + 5GHz\n802.11ax Wi-Fi 6', '~£35–45', 'Overkill — ESP32-S3 only\ndoes 802.11n/b/g 2.4GHz.\nExtra cost, no benefit for CSI'],
], ['Device', 'Mode', 'WiFi Standard', 'Price', 'Verdict'])

add_warning('The TP-Link WA801N is the right buy. Plug it into your room ethernet port via a cable, set it to Access Point mode, give it its own SSID (e.g. "echopose-mesh"). Your ESP32s connect to THIS network only. Your laptop stays on student WiFi for internet.', 'NOTE')

add_warning('Make sure to use a LONG ethernet cable (at least 2m) so you can position the AP centrally in the room for best coverage.', 'TIP')

divider()
add_heading('0.3  ESP32 Nodes', 2, C_BLUE)
add_para('You currently have:', bold=True)
add_hardware_table([
    ['Waveshare ESP32-S3 Dev Board ×2', 'PCB trace antenna (built-in)', 'GOOD — use for Nodes 1 & 2'],
    ['NodeMCU-32S ESP32-32D (original ESP32)', 'PCB trace antenna', 'PROBLEM — original ESP32, NOT S3. CSI on original ESP32 is limited. Use as spare/test only.'],
    ['1× more ESP32-S3 needed', 'MUST have U.FL external antenna connector', 'Buy this — see below'],
], ['Device You Have', 'Antenna Type', 'Status'])

add_warning('The NodeMCU-32S is the original ESP32 (Xtensa LX6), NOT ESP32-S3 (Xtensa LX7). CSI capture on original ESP32 is unstable and limited to fewer subcarriers. Do NOT use it as a primary data node. Your firmware targets ESP32-S3.', 'CRITICAL')

add_para('For Node 3 — buy one of these (in order of preference):', bold=True)
add_hardware_table([
    ['Waveshare ESP32-S3-Zero\n(with U.FL connector)', 'ESP32-S3-WROOM-1 module\n+ U.FL external antenna socket', '~£8–12', 'Best — compact, has external antenna port, same chip as your other 2'],
    ['ESP32-S3-DevKitC-1\n(U.FL variant)', 'ESP32-S3 + U.FL socket', '~£10–15', 'Good — official Espressif devkit with antenna option'],
    ['Waveshare ESP32-S3 Dev Board\n(same as yours)', 'PCB trace antenna only', '~£12', 'Acceptable if U.FL not available — same as your current 2 nodes'],
], ['Board', 'Antenna', 'Price', 'Notes'])

add_warning('When buying, look specifically for "U.FL" or "IPEX" or "external antenna" in the product title. Then buy a 2.4GHz WiFi antenna (~£3) with a U.FL to SMA connector. This gives you +3–6 dBi signal strength and better coverage.', 'TIP')

divider()
add_heading('0.4  Complete Shopping List', 2, C_BLUE)
add_hardware_table([
    ['Xbox One Kinect 2 Sensor',    'CeX',         'MUST BUY',   '~£25–40'],
    ['Kinect Adapter for Windows',  'CeX / eBay',  'MUST BUY',   '~£15–25'],
    ['TP-Link TL-WA801N AP',        'Amazon / eBay','MUST BUY',  '~£20'],
    ['Ethernet cable (2m+)',        'Anywhere',    'MUST BUY',   '~£4'],
    ['ESP32-S3 (U.FL antenna)',     'Amazon / AliExpress','MUST BUY','~£8–15'],
    ['2.4GHz external WiFi antenna','Amazon',      'RECOMMENDED','~£3–5'],
    ['U.FL to SMA adapter cable',   'Amazon',      'RECOMMENDED','~£3'],
    ['USB 3.0 hub (powered)',       'Amazon',      'RECOMMENDED','~£12–20'],
    ['Floor mat / crash mat',       'Sports Direct','FOR SAFETY','~£15–25'],
    ['Measuring tape (5m)',         'Pound shop',  'FOR SETUP',  '~£2'],
], ['Item', 'Where to Buy', 'Priority', 'Approx. Cost'])

add_para('Estimated total outlay: £110 – £160', bold=True, color=C_GREEN, size=12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 1 — PHYSICAL SETUP
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 1   PHYSICAL SETUP — NODES, ROUTER & CAMERA', C_TEAL, C_WHITE, 14)
doc.add_paragraph()

add_heading('1.1  Room Layout', 2, C_TEAL)
add_ascii_box('Ideal Room Layout (Top-Down View)', [
    '  ┌────────────────────────────────────────────────────────────────┐',
    '  │                        YOUR ROOM                              │',
    '  │                                                                │',
    '  │  [NODE-0]                              [NODE-2]               │',
    '  │  ESP32-S3                              ESP32-S3               │',
    '  │  top-left corner                       top-right corner       │',
    '  │  height: 1.8m                          height: 1.8m           │',
    '  │                                                                │',
    '  │                  ★ DATA COLLECTION ZONE ★                     │',
    '  │                  (3m × 3m clear floor space)                  │',
    '  │                  Person walks / performs here                  │',
    '  │                                                                │',
    '  │  [NODE-1]                   [KINECT + LAPTOP]                 │',
    '  │  ESP32-S3                   at far wall, 1.0m height          │',
    '  │  bottom-centre              faces into room                    │',
    '  │  height: 0.3m               USB3 → laptop                    │',
    '  │  (low, floor-level)                                            │',
    '  │                                                                │',
    '  │  [TP-Link AP]  ──ethernet──  [Room Ethernet Port]             │',
    '  │  mounted high on wall                                          │',
    '  └────────────────────────────────────────────────────────────────┘',
    '',
    '  Node positions form a TRIANGLE around the collection zone.',
    '  Nodes at different heights = better 3D reconstruction.',
    '  Kinect faces the zone from the SAME side as the laptop.',
], C_TEAL)

add_heading('1.2  Node Placement Rules', 2, C_TEAL)
add_hardware_table([
    ['Node 0', 'Top-left corner', '1.8m (above head)', 'Angled 45° toward centre', 'Line of sight to zone'],
    ['Node 1', 'Bottom wall centre', '0.3m (near floor)', 'Level / slightly up', 'Captures low-body CSI'],
    ['Node 2', 'Top-right corner', '1.8m (above head)', 'Angled 45° toward centre', 'Line of sight to zone'],
], ['Node', 'Position', 'Height', 'Angle', 'Requirement'])

add_warning('Never move the nodes between data collection sessions. The model learns the specific geometry of your setup. If you move a node even slightly, the collected data becomes incompatible.', 'CRITICAL')
add_warning('Mark node positions with tape on the floor/wall. Photograph the setup before and after every session.', 'TIP')

divider()
add_heading('1.3  Router / AP Setup Steps', 2, C_TEAL)

add_step_box('1', 'Plug TP-Link WA801N into room ethernet port', [
    '   • Connect ethernet cable: wall socket → WA801N WAN port',
    '   • Power the WA801N via USB or included PSU',
    '   • Wait 60 seconds for it to boot',
], C_TEAL)

add_step_box('2', 'Configure WA801N as Access Point', [
    '   • Connect laptop to WA801N via ethernet or default WiFi',
    '   • Open browser → http://192.168.0.1',
    '   • Login: admin / admin (default)',
    '   • Go to: Quick Setup → Access Point mode',
    '   • Set SSID: echopose-mesh',
    '   • Set Password: echopose2025 (or your choice)',
    '   • Channel: 6 (fixed, do NOT set auto)',
    '   • Bandwidth: 20MHz (NOT 40MHz — for stable CSI)',
    '   • Save & reboot',
], C_TEAL)

add_step_box('3', 'Update ESP32 firmware WiFi credentials', [
    '   • Open: firmware/sdkconfig.defaults',
    '   • Change: CONFIG_WIFI_SSID="echopose-mesh"',
    '   • Change: CONFIG_WIFI_PASSWORD="echopose2025"',
    '   • Change: CONFIG_HOST_IP=<your laptop IP on that AP network>',
    '   • Flash each ESP32 with idf.py flash',
], C_TEAL)

add_warning('Set the AP channel to a FIXED number (6 is ideal for 2.4GHz). If the channel changes, your CSI readings change character. This breaks cross-session consistency.', 'WARNING')

divider()
add_heading('1.4  Kinect Setup Steps', 2, C_TEAL)

add_step_box('1', 'Connect Kinect 2 to laptop via Kinect Adapter', [
    '   • Kinect 2 cable → Kinect Adapter box',
    '   • Kinect Adapter USB 3.0 → laptop USB 3.0 port (blue port)',
    '   • Kinect Adapter power cable → wall socket',
    '   • IMPORTANT: USB 3.0 is required — USB 2.0 will NOT work',
], C_RED)

add_step_box('2', 'Install Kinect SDK (Windows)', [
    '   • Download: "Kinect for Windows SDK 2.0" from Microsoft',
    '   • URL: microsoft.com/en-us/download/details.aspx?id=44561',
    '   • Run installer, follow prompts, restart laptop',
    '   • Open "Kinect Studio" — verify depth view appears',
], C_TEAL)

add_step_box('3', 'Install Python Kinect bridge', [
    '   • pip install pykinect2',
    '   • This gives Python access to Kinect skeleton data',
    '   • Test: python scripts/test_kinect.py',
], C_TEAL)

add_warning('If your laptop does NOT have USB 3.0, you cannot use Kinect 2. In that case, fall back to: Brio 500 + MediaPipe (2D poses only). This still works but gives less accurate 3D ground truth.', 'WARNING')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 2 — SOFTWARE SETUP
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 2   SOFTWARE SETUP — PC STACK INSTALLATION', C_BLUE, C_WHITE, 14)
doc.add_paragraph()

add_heading('2.1  Prerequisites — Install These First', 2, C_BLUE)
add_hardware_table([
    ['Python 3.11+',    'python.org/downloads',                    'pip install -r inference/requirements.txt'],
    ['Rust + Cargo',    'rustup.rs',                               'curl --proto https --tlsv1.2 -sSf https://sh.rustup.rs | sh'],
    ['ESP-IDF v5.x',    'docs.espressif.com → Get Started',       'Windows installer: dl.espressif.com/dl/esp-idf'],
    ['Docker Desktop',  'docker.com/products/docker-desktop',     'For production deploy (optional for dev)'],
    ['Git',             'git-scm.com',                            'git --version to verify'],
    ['VS Code',         'code.visualstudio.com',                  '+ Rust Analyzer + Python extensions'],
], ['Software', 'Where to Get', 'Install Command / Notes'])

divider()
add_heading('2.2  Clone & Install (One-Time)', 2, C_BLUE)

add_ascii_box('Terminal Commands — Run In Order', [
    '  # 1. Navigate to your project',
    '  cd "C:\\Users\\Admin\\wifi vision"',
    '',
    '  # 2. Install Python dependencies',
    '  pip install -r inference/requirements.txt',
    '',
    '  # 3. Build Rust aggregator',
    '  cd aggregator',
    '  cargo build --release',
    '  cd ..',
    '',
    '  # 4. Verify everything works with mock data',
    '  # Terminal 1:',
    '  cargo run --release --manifest-path aggregator/Cargo.toml',
    '',
    '  # Terminal 2:',
    '  python inference/server.py',
    '',
    '  # Terminal 3:',
    '  python scripts/mock_esp32_mesh.py',
    '',
    '  # Browser:',
    '  Open ui/index.html — you should see animated skeleton',
], C_BLUE)

add_warning('Run the mock first before touching real hardware. If the mock works, your stack is healthy. If the mock fails, fix it before connecting ESP32 boards.', 'TIP')

divider()
add_heading('2.3  Flash ESP32-S3 Nodes', 2, C_BLUE)

add_step_box('1', 'Set up ESP-IDF environment', [
    '   Windows: Open "ESP-IDF Command Prompt" (installed with ESP-IDF)',
    '   Linux/Mac: . $HOME/esp/esp-idf/export.sh',
], C_BLUE)

add_step_box('2', 'Configure each node', [
    '   cd firmware',
    '   idf.py menuconfig',
    '   → Component config → EchoPose Config',
    '   → Set NODE_ID = 0 (for first board)',
    '   → Set WIFI_SSID = echopose-mesh',
    '   → Set HOST_IP = <your laptop IP>',
    '   Save and exit',
], C_BLUE)

add_step_box('3', 'Flash and monitor', [
    '   Connect ESP32-S3 via USB-C',
    '   idf.py -p COM3 flash monitor   (change COM3 to your port)',
    '   You should see: "CSI capture started, streaming to HOST_IP:5005"',
    '   Repeat steps 2-3 for NODE_ID = 1 and NODE_ID = 2',
], C_BLUE)

add_step_box('4', 'Verify all 3 nodes are seen by aggregator', [
    '   GET http://localhost:3000/health',
    '   Response should show: "nodes_connected": 3',
    '   If less than 3, check WiFi connection on missing node',
], C_BLUE)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 3 — DATA COLLECTION
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 3   DATA COLLECTION — RECORDING GROUND TRUTH', C_ORANGE, C_WHITE, 14)
doc.add_paragraph()

add_heading('3.1  How Data Collection Works', 2, C_ORANGE)

add_ascii_box('Two Streams — Synchronized By Timestamp', [
    '                  ┌─────────────────────────────┐',
    '                  │        YOUR ROOM             │',
    '  ┌──────────┐    │  Person walking, sitting,   │    ┌──────────────┐',
    '  │ ESP32 ×3 │    │  falling, waving, etc.      │    │ Kinect 2 /  │',
    '  │          │    │                              │    │ Brio 500 +  │',
    '  │ CSI at   │    │  [WiFi signals bounce       │    │ MediaPipe   │',
    '  │ 20Hz     │    │   off person]               │    │             │',
    '  └────┬─────┘    └─────────────────────────────┘    └──────┬───────┘',
    '       │ UDP frames                                          │ Skeleton',
    '       │ (timestamped)                                       │ frames',
    '       ▼                                                     │ (timestamped)',
    '  ┌──────────┐                                               │',
    '  │ Aggreg-  │                                               │',
    '  │ ator     │                                               │',
    '  └────┬─────┘                                               │',
    '       │                                                     │',
    '       ▼                                                     ▼',
    '  ┌──────────────────────────────────────────────────────────────┐',
    '  │              DATA COLLECTOR SCRIPT                           │',
    '  │  Aligns CSI windows ↔ Skeleton poses by timestamp           │',
    '  │  Saves: sessions/session_001.npz                             │',
    '  │  Each file: (CSI_windows, pose_labels, timestamps)           │',
    '  └──────────────────────────────────────────────────────────────┘',
], C_ORANGE)

divider()
add_heading('3.2  Data Collection Protocol', 2, C_ORANGE)
add_para('Each session = 1 person, 1 room setup, ~20–30 minutes of activity.', bold=True)
add_para('You need: MINIMUM 5 people × 3 different room arrangements = 15 sessions.')
doc.add_paragraph()

add_hardware_table([
    ['Activity', 'Duration Each', 'Why Important'],
    ['Walking across room (left↔right)', '3 minutes', 'Most common pose, full gait cycle'],
    ['Walking toward/away from Kinect', '3 minutes', 'Depth variation, occlusion handling'],
    ['Standing still', '2 minutes', 'Baseline / stillness calibration'],
    ['Sitting down and standing up', '3 minutes', 'Transition detection, fall risk'],
    ['Raising both arms above head', '2 minutes', 'Upper body keypoint training'],
    ['Waving (left hand, right hand)', '2 minutes', 'Gesture recognition ground truth'],
    ['Controlled fall (onto crash mat)', '5 falls total', 'CRITICAL for fall detection accuracy'],
    ['Lying down (on mat)', '2 minutes', 'Post-fall / sleep position'],
    ['Two people in frame at once', '3 minutes', 'Multi-person tracking ground truth'],
], ['Activity', 'Duration', 'Why Important'])

add_warning('For falls: USE A CRASH MAT. Do controlled slow falls — sit then roll sideways. Never do a real fall. You need the CSI pattern of a falling person, not an injured one.', 'WARNING')

divider()
add_heading('3.3  Running the Data Collector', 2, C_ORANGE)

add_ascii_box('Terminal Commands — Data Collection', [
    '  # Terminal 1: Start aggregator (always running)',
    '  cargo run --release --manifest-path aggregator/Cargo.toml',
    '',
    '  # Terminal 2: Start data collector',
    '  python scripts/collect_data.py \\',
    '      --session-id 001 \\',
    '      --subject "person_name" \\',
    '      --room "bedroom" \\',
    '      --camera kinect2',
    '      # OR: --camera mediapipe  (if using Brio 500)',
    '',
    '  # The script will say:',
    '  # "CSI stream connected: 3 nodes"',
    '  # "Camera stream connected: 30fps"',
    '  # "READY — Press ENTER to start recording"',
    '',
    '  # Before pressing ENTER:',
    '  # 1. Clap your hands ONCE in front of camera',
    '  #    (this creates a sync spike)',
    '  # 2. Then press ENTER',
    '  # 3. Start doing activities',
    '',
    '  # Press Q to stop. File saved to: data/sessions/session_001.npz',
], C_ORANGE)

add_step_box('!', 'CLAP SYNC — Critical Step', [
    '   Before EVERY session:',
    '   1. Stand in front of the Kinect/camera',
    '   2. Clap your hands loudly ONCE',
    '   3. Wait 2 seconds',
    '   4. Then start activities',
    '',
    '   The clap creates a spike in BOTH the CSI signal AND the camera feed.',
    '   The data collector uses this spike to align the two streams perfectly.',
    '   Without this, poses will be off by several frames and training will fail.',
], C_RED)

divider()
add_heading('3.4  Data Quality Checklist', 2, C_ORANGE)
add_para('After each session, verify these before moving on:', bold=True)
add_checklist([
    'Session file saved: data/sessions/session_XXX.npz',
    'File size > 50MB (if smaller, data may be incomplete)',
    'Run: python scripts/verify_session.py --session 001',
    'Sync offset < 50ms (script will report this)',
    'At least 1 fall in the session',
    'At least 5 minutes of walking',
    'Kinect skeleton was visible throughout (check script output)',
    'No node disconnected mid-session (check aggregator logs)',
])

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 4 — MODEL TRAINING
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 4   MODEL TRAINING — RUNNING & EVALUATING', C_GREEN, C_WHITE, 14)
doc.add_paragraph()

add_heading('4.1  Training Pipeline Overview', 2, C_GREEN)

add_ascii_box('Data Flow: Raw Sessions → Trained Model', [
    '  data/sessions/                   Preprocessing',
    '  ├── session_001.npz  ──────────► extract_windows()',
    '  ├── session_002.npz  ──────────► denoise_pipeline()',
    '  ├── session_003.npz  ──────────► fft_doppler()',
    '  └── ...                         normalize()',
    '                                      │',
    '                                      ▼',
    '                               data/processed/',
    '                               ├── train_set.pt   (70%)',
    '                               ├── val_set.pt     (15%)',
    '                               └── test_set.pt    (15%)',
    '                                      │',
    '                                      ▼',
    '                               python train.py',
    '                               ├── PoseNetV2 model',
    '                               ├── Optimizer: AdamW',
    '                               ├── Loss: MSE + confidence BCE',
    '                               ├── Epochs: 100',
    '                               └── Checkpoint: best_model.pt',
    '                                      │',
    '                                      ▼',
    '                               Evaluation',
    '                               ├── MPJPE (target: < 0.15m)',
    '                               ├── PCK@0.1 (target: > 70%)',
    '                               └── Fall Detection F1 (target: > 0.85)',
], C_GREEN)

divider()
add_heading('4.2  Preparing the Data', 2, C_GREEN)

add_ascii_box('Terminal — Preprocess All Sessions', [
    '  # Split by ROOM not by time',
    '  # This tests if model generalises to new environments',
    '',
    '  python scripts/preprocess_data.py \\',
    '      --sessions-dir data/sessions/ \\',
    '      --output-dir data/processed/ \\',
    '      --train-rooms bedroom1,lounge,kitchen \\',
    '      --val-rooms  bedroom2 \\',
    '      --test-rooms bedroom3',
    '',
    '  # This runs your existing denoising pipeline on each window:',
    '  # 1. Adaptive Wiener filtering',
    '  # 2. Wavelet denoising (db4)',
    '  # 3. Spectral subtraction',
    '  # 4. FFT Doppler spectrum extraction',
    '  # Output: numpy arrays ready for the model',
], C_GREEN)

add_warning('ALWAYS split by room, not by time. If you split by time, your training and test data are from the SAME room, and the model just memorises the room layout. It will fail in a customer\'s house.', 'CRITICAL')

divider()
add_heading('4.3  Running Training', 2, C_GREEN)

add_ascii_box('Terminal — Training Command', [
    '  python scripts/train_model.py \\',
    '      --data-dir data/processed/ \\',
    '      --checkpoint-dir inference/models/ \\',
    '      --epochs 100 \\',
    '      --batch-size 32 \\',
    '      --lr 1e-3 \\',
    '      --device auto   # auto = GPU if available, else CPU',
    '',
    '  # Expected output per epoch:',
    '  # Epoch 01/100 | Train Loss: 2.341 | Val MPJPE: 0.412m | PCK: 18%',
    '  # Epoch 10/100 | Train Loss: 1.102 | Val MPJPE: 0.298m | PCK: 41%',
    '  # Epoch 50/100 | Train Loss: 0.421 | Val MPJPE: 0.198m | PCK: 62%',
    '  # Epoch 100/100| Train Loss: 0.198 | Val MPJPE: 0.142m | PCK: 74%',
    '  # Best model saved: inference/models/pose_net.pt',
    '',
    '  # On your laptop CPU: ~8–16 hours for 100 epochs',
    '  # On Google Colab GPU (free): ~2–4 hours',
    '  # On a £300 RTX 3060 GPU: ~45 minutes',
], C_GREEN)

add_warning('Use Google Colab (colab.research.google.com) for free GPU access. Upload your data/processed/ folder to Google Drive, mount it in Colab, and run training there. Free T4 GPU = 10x faster than your laptop CPU.', 'TIP')

divider()
add_heading('4.4  Is the Model Good Enough?', 2, C_GREEN)
add_hardware_table([
    ['MPJPE (mean error)', 'Mean distance from predicted joint to real joint', 'Not usable', 'Acceptable', 'Good', 'Excellent'],
    ['PCK@0.1 (accuracy)', 'Joints within 10% of person height', '< 30%', '30–50%', '50–70%', '> 70%'],
    ['Fall Detection F1', 'Precision × Recall for fall events', '< 0.70', '0.70–0.80', '0.80–0.90', '> 0.90'],
    ['Latency (inference)', 'Time from CSI window to skeleton output', '> 100ms', '50–100ms', '25–50ms', '< 25ms'],
], ['Metric', 'What It Means', '🔴 Bad', '🟡 OK', '🟢 Good', '✅ Ship It'])

add_warning('Do NOT ship the product until PCK@0.1 > 50% and Fall F1 > 0.80. Below these thresholds, customers will see wrong keypoints and missed falls — they will ask for refunds.', 'CRITICAL')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 5 — PRODUCT BUILD
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 5   PRODUCT BUILD — ALERTS, DASHBOARD & INSTALLER', C_NAVY, C_WHITE, 14)
doc.add_paragraph()

add_heading('5.1  Alert System', 2, C_NAVY)
add_para('Your FallDetector and HealthAnomalyDetector already work. You need to wire them to actual notifications.', italic=True)
doc.add_paragraph()

add_hardware_table([
    ['Fall Detected',    'IMMEDIATE',  'Email + SMS + Push notification', 'FallDetector.risk_score > 0.8'],
    ['Person Not Moving\n> 5 min',  'HIGH',  'Email + Push', 'OccupancyAnalyzer + stillness timer'],
    ['Vital Anomaly',   'MEDIUM',     'Email',          'HealthAnomalyDetector → CRITICAL status'],
    ['Intruder\n(unexpected person)', 'HIGH', 'SMS + Email', 'Occupancy when system armed'],
    ['Node Offline',    'SYSTEM',     'Email',          'Aggregator health check fails'],
    ['Daily Summary',   'LOW',        'Email',          '23:00 daily digest (configurable)'],
], ['Alert Type', 'Urgency', 'Channels', 'Trigger Condition'])

add_ascii_box('Alert System Architecture', [
    '  inference/pipeline/                  notification/',
    '  ├── fall_detector.py  ─── event ──► alert_manager.py',
    '  ├── health_anomaly.py ─── event ──► │',
    '  ├── occupancy.py      ─── event ──► │  ┌─────────────┐',
    '  └── tactical/...      ─── event ──► └──► EmailSender  │ (SendGrid free tier)',
    '                                          │ SmsSender   │ (Twilio £10 credit)',
    '                                          │ PushSender  │ (Firebase free)',
    '                                          │ WebhookSend │ (HTTP POST to URL)',
    '                                          └─────────────┘',
    '',
    '  Customer configures in dashboard:',
    '  "Send me SMS when a fall is detected"',
    '  "Email me daily summary at 8am"',
    '  "POST to my Home Assistant webhook URL"',
], C_NAVY)

divider()
add_heading('5.2  Customer Dashboard (Replace Dev UI)', 2, C_NAVY)
add_para('The current ui/ folder is a developer debug view. Replace with a proper product UI.', italic=True)
doc.add_paragraph()

add_hardware_table([
    ['Login / Account',    'Week 1', 'Email + password, JWT tokens, "Remember me"'],
    ['Room Setup Wizard',  'Week 1', 'Step-by-step: "Place node 1 here" → photo guide → verify'],
    ['Live Monitor View',  'Week 2', 'Simplified 3D skeleton + status badges (NORMAL / ALERT)'],
    ['Alert History',      'Week 2', 'List of all past alerts, timestamp, severity, dismissed/actioned'],
    ['Settings Panel',     'Week 2', 'Notification preferences, alert thresholds, account details'],
    ['Analytics Reports',  'Week 3', 'Daily/weekly graphs: activity levels, sleep quality, fall incidents'],
    ['Hardware Status',    'Week 3', 'Node battery/signal, last-seen timestamp, calibration button'],
    ['User Management',    'Week 3', 'For care facilities: multiple rooms, multiple staff accounts'],
], ['Screen / Feature', 'Build Week', 'What It Contains'])

divider()
add_heading('5.3  Single-Click Installer', 2, C_NAVY)

add_ascii_box('Installer Architecture (Windows .exe)', [
    '  EchoPose_Installer_v1.0.exe',
    '  │',
    '  ├── Bundled inside:',
    '  │   ├── Docker Desktop (auto-installs if missing)',
    '  │   ├── echopose-aggregator Docker image',
    '  │   ├── echopose-inference Docker image',
    '  │   ├── nginx-gateway Docker image',
    '  │   └── Startup scripts',
    '  │',
    '  └── Installer steps (what customer sees):',
    '      [1/5] Checking system requirements...',
    '      [2/5] Installing EchoPose engine...',
    '      [3/5] Configuring network...',
    '      [4/5] Setting up dashboard...',
    '      [5/5] Done! Opening EchoPose at http://localhost:8000',
    '',
    '  Tool to build this: Inno Setup (free) — wraps Docker Compose',
    '  Or: Electron app shell with embedded Docker management',
], C_NAVY)

add_warning('Use Inno Setup (free, jrsoftware.org/isinfo.php) to create the .exe installer. It bundles your Docker images as offline files so the customer does not need internet to install.', 'TIP')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  PART 6 — DEPLOYMENT & MONETIZATION
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  PART 6   DEPLOYMENT — SAAS & ON-PREMISE + MONETIZATION', C_TEAL, C_WHITE, 14)
doc.add_paragraph()

add_heading('6.1  Two Products, One Codebase', 2, C_TEAL)

add_ascii_box('Product Tiers', [
    '  ┌─────────────────────────────────┐  ┌───────────────────────────────────┐',
    '  │     ECHOPOSE CLOUD (SaaS)       │  │  ECHOPOSE ON-PREMISE (License)    │',
    '  │                                 │  │                                   │',
    '  │  Customer buys monthly sub      │  │  Customer pays once               │',
    '  │  You host the inference server  │  │  They run it on their server      │',
    '  │  They just install the local    │  │  You give them a license key      │',
    '  │  "edge agent" (small app)       │  │                                   │',
    '  │                                 │  │  TARGET: Hospitals, care homes,   │',
    '  │  TARGET: Homeowners, families,  │  │  security firms, enterprises      │',
    '  │  small care agencies            │  │                                   │',
    '  │                                 │  │  Price: £2,000–10,000 one-time   │',
    '  │  Price: £49–199/month           │  │  + £500/yr support optional       │',
    '  │                                 │  │                                   │',
    '  │  Your infrastructure cost:      │  │  Your cost after sale: ~£0        │',
    '  │  ~£50–100/month (VPS)           │  │  (they run it themselves)         │',
    '  └─────────────────────────────────┘  └───────────────────────────────────┘',
], C_TEAL)

divider()
add_heading('6.2  Monetization Layer — What to Build', 2, C_TEAL)

add_hardware_table([
    ['Stripe account',           'stripe.com',         'Free',        'Payment processing for SaaS subscriptions'],
    ['Stripe Billing',           'In your dashboard',  'Free + fees', 'Monthly subscriptions, usage metering'],
    ['License key generator',    'Build it (~1 day)',  '£0',          'HMAC-SHA256 signed key: customer_id+expiry+features'],
    ['License validator',        'In aggregator/inference', '£0',     'On-premise: checks key on startup, phones home monthly'],
    ['Customer portal',          'Stripe Customer Portal','Free',     'Customers manage their own subscription, invoices'],
    ['Twilio (SMS alerts)',       'twilio.com',         '£10 credit',  'For fall/emergency SMS notifications'],
    ['SendGrid (email)',          'sendgrid.com',       'Free tier',   '100 emails/day free — enough for MVP'],
], ['Component', 'Where / Tool', 'Cost', 'Purpose'])

divider()
add_heading('6.3  Go-Live Checklist', 2, C_TEAL)
add_para('Before you take your first paying customer:', bold=True, color=C_RED)
add_checklist([
    'PCK@0.1 accuracy > 50% on test rooms you have never trained on',
    'Fall detection F1 > 0.80 (test with 20+ controlled falls)',
    'Alert system tested: email, SMS both delivered in < 30 seconds',
    'Installer tested on a CLEAN laptop that has never had your code',
    'Privacy policy written (you MUST have one — you sense people)',
    'Terms of service written (liability for missed fall alerts)',
    'Stripe live mode enabled and test transaction processed',
    'Customer dashboard login working with real account',
    'At least 1 beta customer (friend/family) using it for 2 weeks',
    'Monitoring in place (you get alerted if your servers go down)',
])

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX A — FULL WIRING DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  APPENDIX A   FULL SYSTEM WIRING DIAGRAM', C_DKGREY, C_WHITE, 13)
doc.add_paragraph()

add_ascii_box('Complete EchoPose Hardware Wiring (Physical)', [
    '  WALL ETHERNET PORT',
    '        │',
    '        │ Cat5e/Cat6 ethernet cable (2m+)',
    '        │',
    '        ▼',
    '  ┌─────────────┐',
    '  │ TP-Link     │  WiFi SSID: "echopose-mesh"',
    '  │ WA801N AP   │  Channel: 6 (fixed)',
    '  │             │  Bandwidth: 20MHz',
    '  └─────────────┘',
    '        │ WiFi 2.4GHz (802.11n)',
    '        │_____________________________________',
    '        │               │                    │',
    '        ▼               ▼                    ▼',
    '  ┌──────────┐   ┌──────────┐        ┌──────────┐',
    '  │ ESP32-S3 │   │ ESP32-S3 │        │ ESP32-S3 │',
    '  │ NODE-0   │   │ NODE-1   │        │ NODE-2   │',
    '  │ top-left │   │ bottom   │        │ top-right│',
    '  │ 1.8m high│   │ 0.3m high│        │ 1.8m high│',
    '  │ USB power│   │ USB power│        │ USB power│',
    '  └──────────┘   └──────────┘        └──────────┘',
    '        │               │                    │',
    '        └───────────────┴────────────────────┘',
    '                        │ UDP port 5005 (WiFi)',
    '                        ▼',
    '  ┌─────────────────────────────────────────────────┐',
    '  │              XIAOMI NOTEBOOK ULTRA               │',
    '  │                                                   │',
    '  │  ┌─────────────┐    ┌──────────────────────┐    │',
    '  │  │   Kinect 2  │    │  EchoPose Stack      │    │',
    '  │  │   Adapter   │    │  ┌──────────────────┐│    │',
    '  │  │   (USB 3.0) │    │  │ Rust Aggregator  ││    │',
    '  │  └──────┬──────┘    │  │ :3000            ││    │',
    '  │         │           │  └────────┬─────────┘│    │',
    '  │         ▼           │           │           │    │',
    '  │  ┌──────────────┐  │  ┌────────▼─────────┐│    │',
    '  │  │  Kinect 2    │  │  │ Python Inference  ││    │',
    '  │  │  Sensor      │  │  │ :8765             ││    │',
    '  │  │  (USB-C/prop)│  │  └────────┬─────────┘│    │',
    '  │  └──────────────┘  │           │           │    │',
    '  │  Brio 500 (USB-A)  │  ┌────────▼─────────┐│    │',
    '  │  ↓ backup camera   │  │ Dashboard        ││    │',
    '  │                    │  │ localhost:8000   ││    │',
    '  │                    │  └──────────────────┘│    │',
    '  │                    └──────────────────────┘    │',
    '  └─────────────────────────────────────────────────┘',
    '                        │',
    '                        │ HTTPS (internet)',
    '                        ▼',
    '             [Stripe / SendGrid / Twilio]',
    '             (billing, email alerts, SMS)',
], C_DKGREY)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX B — ENV VARS
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  APPENDIX B   ENVIRONMENT VARIABLE REFERENCE', C_DKGREY, C_WHITE, 13)
doc.add_paragraph()
add_hardware_table([
    ['AGGREGATOR_UDP_PORT',   '5005',              'UDP port ESP32 nodes send to'],
    ['AGGREGATOR_HTTP_PORT',  '3000',              'Aggregator HTTP + WebSocket port'],
    ['AGGREGATOR_SYNC_WINDOW_MS','50',             'How long to wait before flushing node bundle'],
    ['EXPECTED_NODES',        '3',                 'How many ESP32 nodes expected'],
    ['INFERENCE_WS_PORT',     '8765',              'Python inference WebSocket port'],
    ['AGGREGATOR_WS_URI',     'ws://localhost:3000/ws','Where inference connects to aggregator'],
    ['INFERENCE_DEVICE',      'auto',              'cpu | cuda | mps | auto'],
    ['ECHOPOSE_API_TOKEN',    '<your-secret-key>', 'API key for /ingest endpoint'],
    ['ECHOPOSE_SESSION_KEY',  '<fernet-key>',      'AES-256 key for session encryption'],
    ['ALLOWED_ORIGINS',       'http://localhost:8000','CORS whitelist for dashboard'],
    ['STRIPE_SECRET_KEY',     'sk_live_...',       'Stripe secret key for billing'],
    ['SENDGRID_API_KEY',      'SG.xxx',            'SendGrid key for email alerts'],
    ['TWILIO_ACCOUNT_SID',    'ACxxx',             'Twilio SID for SMS alerts'],
    ['TWILIO_AUTH_TOKEN',     'xxx',               'Twilio auth token'],
    ['ALERT_EMAIL',           'customer@email.com','Default alert recipient'],
    ['ALERT_PHONE',           '+447700000000',     'Default SMS recipient'],
], ['Variable', 'Default / Example', 'Purpose'])

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX C — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

add_coloured_para('  APPENDIX C   TROUBLESHOOTING CHEAT SHEET', C_DKGREY, C_WHITE, 13)
doc.add_paragraph()

add_hardware_table([
    ['ESP32 not connecting to AP',
     'Wrong SSID/password in sdkconfig\nAP on wrong channel\nAP not in AP mode',
     'Check sdkconfig.defaults\nVerify AP SSID in WA801N admin\nPing ESP32 IP from laptop'],
    ['Aggregator shows < 3 nodes',
     'Node not powered\nNode firmware not flashed\nFirewall blocking UDP 5005',
     'Check Windows Firewall → allow UDP 5005\nCheck each ESP32 serial monitor'],
    ['Kinect not detected in Device Manager',
     'No Kinect Adapter\nUsing USB 2.0 port\nDriver not installed',
     'MUST use USB 3.0 (blue port)\nInstall Kinect SDK 2.0\nCheck Device Manager for errors'],
    ['Model outputs wrong keypoints',
     'Current model trained on synthetic data\nNot enough real training data',
     'Collect real sessions\nRetrain model\nCheck PCK metric'],
    ['CSI sync offset > 100ms',
     'Missed clap sync\nClock drift between nodes',
     'Redo session with clear clap\nCheck NTP sync on laptop'],
    ['Inference server crashes',
     'Missing Python dependency\nGPU out of memory',
     'pip install -r requirements.txt\nSet INFERENCE_DEVICE=cpu'],
    ['Dashboard shows no skeleton',
     'Wrong WebSocket port\nBrowser CORS block',
     'Check ALLOWED_ORIGINS env var\nOpen browser console for errors'],
    ['Alert email not arriving',
     'SendGrid key wrong\nIn spam folder\nFree tier limit hit',
     'Check SENDGRID_API_KEY\nVerify sender domain in SendGrid'],
], ['Problem', 'Likely Causes', 'Fix'])

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EchoPose  ·  Confidential & Commercial in Confidence  ·  shazin2889@gmail.com')
run.font.size  = Pt(8)
run.font.color.rgb = C_MGREY

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = 'EchoPose_Guide.docx'
doc.save(output_path)
print(f'Guide saved: {output_path}')
