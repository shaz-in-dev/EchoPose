"""
EchoPose — HR & Legal Documents Generator
Produces: Employee Handbook + Mutual NDA + One-Way NDA
Run: python scripts/generate_hr_legal.py
Output: docs/legal/
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = 'docs/legal'
os.makedirs(OUT, exist_ok=True)

OWNER       = 'Muhammed Shazin Sadhik Kunhi Parambath'
TRADING_AS  = 'EchoPose'
EMAIL       = 'shazin2889@gmail.com'
COUNTRY     = 'England and Wales'

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
BLUE   = RGBColor(0x1B, 0x6C, 0xA8)
TEAL   = RGBColor(0x00, 0x87, 0x8A)
GREEN  = RGBColor(0x1A, 0x73, 0x48)
RED    = RGBColor(0xC0, 0x20, 0x20)
ORANGE = RGBColor(0xD9, 0x6B, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF2, 0xF4, 0xF7)
MGREY  = RGBColor(0x99, 0xA3, 0xB0)

def h(c): return f'{c[0]:02X}{c[1]:02X}{c[2]:02X}'

def shd(cell, rgb):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    e = OxmlElement('w:shd')
    e.set(qn('w:val'),'clear'); e.set(qn('w:color'),'auto')
    e.set(qn('w:fill'), h(rgb)); pr.append(e)

class Doc:
    def __init__(self, filename, accent=BLUE):
        self.d = Document(); self.fn = filename; self.ac = accent
        s = self.d.sections[0]
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin  = s.bottom_margin = Cm(2.5)

    def save(self):
        p = os.path.join(OUT, self.fn); self.d.save(p); print(f'  Saved: {p}')

    def pb(self): self.d.add_page_break()

    def cover(self, title, subtitle, doc_label):
        t = self.d.add_table(1,1)
        c = t.cell(0,0); shd(c, NAVY); c.width = Inches(6)
        for txt,sz,bold,col in [
            (TRADING_AS, 28, True, TEAL),
            (title, 18, True, WHITE),
            ('', 6, False, WHITE),
            (subtitle, 11, False, RGBColor(0xB0,0xD0,0xFF)),
            ('', 6, False, WHITE),
            (doc_label, 9, True, ORANGE),
            (f'{OWNER}  ·  {EMAIL}', 8, False, MGREY),
        ]:
            p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt); r.font.size=Pt(sz)
            r.font.bold=bold; r.font.color.rgb=col
        self.d.add_paragraph(); self.pb()

    def banner(self, text, color=None):
        color = color or self.ac
        t = self.d.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.LEFT
        c = t.cell(0,0); shd(c, color); c.width=Inches(6)
        p = c.paragraphs[0]
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text); r.font.bold=True
        r.font.size=Pt(12); r.font.color.rgb=WHITE
        self.d.add_paragraph()

    def h1(self, txt):
        p = self.d.add_heading(txt, 1)
        for r in p.runs: r.font.color.rgb=self.ac; r.font.bold=True

    def h2(self, txt):
        p = self.d.add_heading(txt, 2)
        for r in p.runs: r.font.color.rgb=self.ac; r.font.bold=True

    def h3(self, txt):
        p = self.d.add_heading(txt, 3)
        for r in p.runs: r.font.color.rgb=NAVY

    def para(self, txt='', bold=False, italic=False, color=None, size=11):
        p = self.d.add_paragraph()
        if txt:
            r = p.add_run(txt); r.font.size=Pt(size)
            r.font.bold=bold; r.font.italic=italic
            if color: r.font.color.rgb=color
        return p

    def bullet(self, items):
        for item in items:
            p = self.d.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(item); r.font.size=Pt(10.5)

    def numbered(self, items):
        for item in items:
            p = self.d.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(item); r.font.size=Pt(10.5)

    def table(self, headers, rows, hcol=None):
        hcol = hcol or self.ac
        t = self.d.add_table(1+len(rows), len(headers))
        t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
        for i,h_ in enumerate(headers):
            c = t.rows[0].cells[i]; shd(c, hcol)
            p = c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h_); r.font.bold=True
            r.font.size=Pt(10); r.font.color.rgb=WHITE
        for ri,row in enumerate(rows):
            bg = LGREY if ri%2==0 else WHITE
            for ci,val in enumerate(row):
                c = t.rows[ri+1].cells[ci]; shd(c, bg)
                p = c.paragraphs[0]; r=p.add_run(str(val)); r.font.size=Pt(9.5)
        self.d.add_paragraph()

    def callout(self, txt, color=None, label=''):
        color = color or self.ac
        t = self.d.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.LEFT
        c = t.cell(0,0); shd(c, color); c.width=Inches(6)
        p = c.paragraphs[0]
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(4)
        if label:
            rl = p.add_run(label+'  '); rl.font.bold=True
            rl.font.size=Pt(10); rl.font.color.rgb=WHITE
        r = p.add_run(txt); r.font.size=Pt(10); r.font.color.rgb=WHITE
        self.d.add_paragraph()

    def sig_block(self, party_name, party_role, show_witness=True):
        self.d.add_paragraph()
        t = self.d.add_table(1, 3 if show_witness else 2)
        t.style = 'Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
        labels = [
            ('Signature', party_name, party_role),
            ('Print Name', '________________________', ''),
            ('Date', '________________________', ''),
        ]
        if show_witness:
            labels += [('Witness Signature', '________________________', 'Witness')]
        # Simpler sig block
        self.para('')
        self.para(f'Signed for and on behalf of {party_name}', bold=True)
        self.para('')
        self.para('Signature:  _______________________________________________')
        self.para('')
        self.para(f'Full Name:  {party_name}')
        self.para('')
        self.para(f'Role / Title:  {party_role}')
        self.para('')
        self.para('Date:  _______________________________________________')
        self.para('')
        if show_witness:
            self.para('Witness Signature:  ____________________________________')
            self.para('Witness Name:  _________________________________________')
            self.para('Witness Address:  ______________________________________')

    def divider(self):
        p = self.d.add_paragraph()
        r = p.add_run('─'*85); r.font.size=Pt(7); r.font.color.rgb=MGREY

    def footer(self, text):
        self.divider()
        p = self.d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size=Pt(8); r.font.color.rgb=MGREY


# ══════════════════════════════════════════════════════════════════════════════
#  EMPLOYEE HANDBOOK
# ══════════════════════════════════════════════════════════════════════════════
print('Building Employee Handbook...')
d = Doc('EchoPose_Employee_Handbook.docx', NAVY)
d.cover(
    'Employee Handbook',
    'Policies · Terms of Employment · Conduct · Rights & Responsibilities',
    'HUMAN RESOURCES  |  Version 1.0'
)

d.callout(
    'This handbook sets out the terms, policies, and expectations that apply to all individuals '
    'engaged by EchoPose (trading name of Muhammed Shazin Sadhik Kunhi Parambath). '
    'It should be read in conjunction with your individual contract of employment or engagement agreement. '
    'This handbook is governed by the laws of England and Wales.',
    NAVY, 'ℹ ABOUT THIS HANDBOOK'
)

d.h1('1.  About EchoPose')
d.h2('1.1  Who We Are')
d.para(
    f'EchoPose is a technology business trading under the name EchoPose, '
    f'owned and operated by {OWNER}. '
    'We build WiFi-based human sensing products that enable real-time fall detection, '
    'vital sign monitoring, and activity tracking — without cameras and without wearables. '
    'Our mission is to make invisible safety technology accessible to every care setting and home.'
)

d.h2('1.2  Our Values')
d.bullet([
    'Privacy first — we build products that protect the dignity of the people they monitor',
    'Quality over speed — we ship when it works, not before',
    'Transparency — we are honest with customers, partners, and each other',
    'Ownership — every team member takes responsibility for their work end to end',
    'Inclusion — we welcome people of all backgrounds, cultures, and identities',
])

d.pb()
d.banner('SECTION 2 — EMPLOYMENT TERMS', BLUE)
d.h1('2.  Terms of Employment')

d.h2('2.1  Contract of Employment')
d.para(
    'Every employee and contractor will be provided with a written statement of employment particulars '
    'within two months of starting, in accordance with the Employment Rights Act 1996 (as amended). '
    'This will include job title, start date, pay, working hours, holiday entitlement, and notice period.'
)

d.h2('2.2  Probationary Period')
d.para(
    'All new employees serve a probationary period of three (3) months. '
    'During this period, the notice period is one week (from either party). '
    'Performance will be reviewed at one month and at the end of the probationary period. '
    'The probationary period may be extended by up to one further month at the company\'s discretion.'
)

d.h2('2.3  Working Hours')
d.table(
    ['Employment Type', 'Standard Hours', 'Flexibility', 'Overtime'],
    [
        ['Full-time employee', '37.5 hours/week', 'Flexible start/end by agreement', 'Paid at standard rate — pre-approved only'],
        ['Part-time employee', 'As per contract', 'Pro-rata of full-time terms', 'Pro-rata'],
        ['Contractor / freelance', 'As per engagement agreement', 'N/A', 'N/A — fixed scope'],
    ]
)
d.para(
    'EchoPose operates primarily remotely. Team members are expected to be available '
    'during core hours of 10:00–16:00 GMT/BST on working days unless otherwise agreed. '
    'Working hours must not exceed 48 hours per week averaged over 17 weeks, '
    'in compliance with the Working Time Regulations 1998, unless a written opt-out is signed.'
)

d.h2('2.4  Salary and Pay')
d.para(
    'Salary is set out in your individual contract. Pay is processed monthly, '
    'on the last working day of each calendar month, directly to your nominated bank account. '
    'EchoPose complies with the National Minimum Wage Act 1998 and National Living Wage requirements. '
    'Pay reviews take place annually, typically in April.'
)

d.h2('2.5  Holiday Entitlement')
d.para(
    'All employees are entitled to 28 days paid holiday per year (inclusive of the 8 UK public/bank holidays), '
    'in accordance with the Working Time Regulations 1998. '
    'Holiday year runs from 1 January to 31 December. '
    'Up to 5 days may be carried over to the following year with prior written approval. '
    'Holiday requests must be submitted at least two weeks in advance.'
)

d.h2('2.6  Sick Leave')
d.para(
    'Employees who are unable to work due to illness or injury must notify their manager '
    'by 9:30am on the first day of absence. '
    'Statutory Sick Pay (SSP) is payable from the fourth qualifying day of absence, '
    'in accordance with current HMRC rules. '
    'A self-certification form is required for absences up to 7 days. '
    'A GP fit note is required for absences of 8 or more days.'
)

d.pb()
d.h2('2.7  Parental Leave')
d.para(
    'EchoPose supports all statutory parental leave rights, including maternity leave (up to 52 weeks), '
    'paternity leave (up to 2 weeks), adoption leave, and shared parental leave, '
    'as set out in the Employment Rights Act 1996 and related regulations. '
    'All parental leave policies are applied equally regardless of gender or family structure.'
)

d.h2('2.8  Notice Periods')
d.table(
    ['Length of Continuous Service', 'Minimum Notice (Employee to Company)', 'Minimum Notice (Company to Employee)'],
    [
        ['During probation (0–3 months)', '1 week', '1 week'],
        ['3 months – 2 years', '1 month', '1 month'],
        ['2 years – 5 years', '1 month', '1 week per year of service (min. 2 weeks)'],
        ['5+ years', '3 months', '1 week per year of service (max. 12 weeks)'],
    ]
)

d.banner('SECTION 3 — CONDUCT & PERFORMANCE', TEAL)
d.h1('3.  Code of Conduct')

d.h2('3.1  General Standards')
d.para(
    'All team members are expected to act professionally, honestly, and with respect '
    'towards colleagues, customers, partners, and the public at all times. '
    'This applies in the workplace, at company events, and in online communications made in a professional capacity.'
)

d.h2('3.2  Confidentiality')
d.para(
    'All employees and contractors must keep confidential all information relating to '
    'EchoPose\'s technology, customers, business strategy, pricing, finances, and trade secrets. '
    'This obligation continues after the end of your employment or engagement. '
    'Breach of confidentiality may result in disciplinary action and/or civil legal proceedings.'
)
d.callout(
    'A separate Non-Disclosure Agreement (NDA) must be signed by all employees, contractors, '
    'and any third party who is given access to EchoPose technology or business information.',
    ORANGE, '⚑ IMPORTANT'
)

d.h2('3.3  Intellectual Property')
d.para(
    'Any work, invention, software, design, or other intellectual property created by an employee '
    'in the course of their employment, or by a contractor under an engagement agreement with an IP assignment clause, '
    'belongs to EchoPose (trading as Muhammed Shazin Sadhik Kunhi Parambath). '
    'This includes work developed on personal equipment if it relates to the company\'s business. '
    'This clause is consistent with the Patents Act 1977 and Copyright, Designs and Patents Act 1988.'
)

d.h2('3.4  Social Media Policy')
d.para(
    'Employees must not post content online that could damage EchoPose\'s reputation, '
    'disclose confidential information, make statements on behalf of EchoPose without authorisation, '
    'or bring the company into disrepute. '
    'Personal views must be clearly identified as personal. '
    'Any work-related social media activity must be pre-approved.'
)

d.h2('3.5  Equality, Diversity and Inclusion')
d.para(
    'EchoPose is committed to providing equal opportunities in employment and will not unlawfully '
    'discriminate on the grounds of age, disability, gender reassignment, marriage or civil partnership, '
    'pregnancy or maternity, race, religion or belief, sex, or sexual orientation, '
    'as protected under the Equality Act 2010.'
)

d.pb()
d.banner('SECTION 4 — DISCIPLINARY & GRIEVANCE', RED)
d.h1('4.  Disciplinary Procedure')
d.para(
    'EchoPose follows the ACAS Code of Practice on Disciplinary and Grievance Procedures. '
    'The following stages apply for misconduct or performance issues:'
)
d.table(
    ['Stage', 'Trigger', 'Action', 'Appeal Right'],
    [
        ['Informal', 'Minor issues, first occurrence', 'Verbal discussion and coaching', 'No formal appeal'],
        ['Stage 1 — Verbal Warning', 'Repeated minor issues or single moderate issue', 'Formal verbal warning, recorded in writing', 'Yes — within 5 working days'],
        ['Stage 2 — Written Warning', 'Further occurrence or more serious issue', 'Formal written warning, 12-month live period', 'Yes — within 5 working days'],
        ['Stage 3 — Final Written Warning', 'Continued issues after written warning', 'Final written warning, 12-month live period', 'Yes — within 5 working days'],
        ['Dismissal', 'Continued misconduct or gross misconduct', 'Employment terminated with notice (or summarily for gross misconduct)', 'Yes — within 10 working days'],
    ]
)
d.para('Examples of gross misconduct (may result in immediate dismissal): theft, fraud, serious breach of confidentiality, harassment, deliberate sabotage of company systems.', italic=True)

d.h1('5.  Grievance Procedure')
d.para(
    'Any employee who has a grievance relating to their employment should raise it as follows:'
)
d.numbered([
    'Raise informally with your manager (or with the owner directly if the grievance involves your manager).',
    'If unresolved within 5 working days, submit a written grievance to Muhammed Shazin Sadhik Kunhi Parambath at ' + EMAIL + '.',
    'A formal meeting will be arranged within 7 working days of receipt.',
    'A written response will be provided within 5 working days of the meeting.',
    'You have the right to appeal the outcome within 5 working days.',
])

d.banner('SECTION 5 — DATA PROTECTION & HEALTH & SAFETY', NAVY)
d.h1('6.  Data Protection')
d.para(
    'EchoPose processes personal data in accordance with UK GDPR and the Data Protection Act 2018. '
    'Employee personal data (name, address, bank details, health information) is held for legitimate employment purposes only. '
    'It will not be shared with third parties without consent, except where required by law. '
    'Employees have the right to access their personal data held by the company. '
    'Any breach of data protection obligations may result in disciplinary action.'
)

d.h1('7.  Health and Safety')
d.para(
    'EchoPose takes its obligations under the Health and Safety at Work etc. Act 1974 seriously. '
    'All team members are responsible for maintaining a safe working environment. '
    'Remote workers are responsible for ensuring their home working environment meets basic DSE (Display Screen Equipment) standards. '
    'Any accident, near-miss, or safety concern must be reported to the owner immediately.'
)

d.h1('8.  Amendments to this Handbook')
d.para(
    'EchoPose reserves the right to amend this handbook at any time. '
    'Employees will be given reasonable notice of any changes. '
    'The most current version of this handbook supersedes all previous versions.'
)

d.para('')
d.para(f'Issued by: {OWNER}', bold=True)
d.para(f'Trading as: {TRADING_AS}')
d.para('Version: 1.0  |  Date: April 2026')
d.para('Governing Law: England and Wales')
d.para('')
d.callout(
    'This handbook is a statement of general policy and does not form part of your contract of employment '
    'unless specifically stated in your individual contract. This document should be reviewed by a qualified '
    'employment solicitor before being issued to any employee.',
    ORANGE, '⚠ LEGAL NOTE'
)
d.footer(f'EchoPose Employee Handbook · {OWNER} · Confidential')
d.save()


# ══════════════════════════════════════════════════════════════════════════════
#  MUTUAL NDA
# ══════════════════════════════════════════════════════════════════════════════
print('Building Mutual NDA...')
d = Doc('EchoPose_Mutual_NDA.docx', NAVY)
d.cover(
    'Mutual Non-Disclosure Agreement',
    'For use when both parties share confidential information',
    'LEGAL TEMPLATE  |  Version 1.0  |  Governing Law: England & Wales'
)

d.callout(
    'Use this agreement when meeting with investors, potential partners, suppliers, or advisors '
    'where information will flow in both directions. Both parties sign.',
    BLUE, 'WHEN TO USE THIS DOCUMENT'
)

d.para('MUTUAL NON-DISCLOSURE AGREEMENT', bold=True, size=14, color=NAVY)
d.para('')

d.para('This Mutual Non-Disclosure Agreement ("Agreement") is entered into as of the date last signed below ("Effective Date") between:', size=11)
d.para('')
d.para(f'Party A:  {OWNER}, trading as {TRADING_AS}', bold=True)
d.para(f'          Contact: {EMAIL}')
d.para('          ("Disclosing Party / Receiving Party")')
d.para('')
d.para('Party B:  ___________________________________________________________', bold=True)
d.para('          Company / Name:  _____________________________________________')
d.para('          Address:  ____________________________________________________')
d.para('          Contact:  ___________________________________________________')
d.para('          ("Disclosing Party / Receiving Party")')
d.para('')
d.para('(Each party may be referred to as a "Party" or collectively as the "Parties".)')
d.para('')

d.h1('1.  Purpose')
d.para(
    'The Parties wish to explore a potential business relationship concerning EchoPose technology, '
    'commercial partnership, investment, and/or other matters of mutual interest '
    '("the Purpose"). In connection with the Purpose, each Party may disclose to the other '
    'certain confidential and proprietary information.'
)

d.h1('2.  Definition of Confidential Information')
d.para(
    '"Confidential Information" means any information disclosed by one Party ("Disclosing Party") '
    'to the other Party ("Receiving Party"), whether orally, in writing, electronically, or by any other means, '
    'that is designated as confidential or that reasonably should be understood to be confidential given the nature '
    'of the information and the circumstances of disclosure. Confidential Information includes, without limitation:'
)
d.bullet([
    'Technical data, trade secrets, know-how, research, product plans, software, and source code',
    'Financial information, pricing, business plans, strategies, forecasts, and projections',
    'Customer lists, customer data, supplier information, and partnership details',
    'Inventions, developments, algorithms, models, and intellectual property of any kind',
    'Any information concerning the EchoPose technology, infrastructure, or architecture',
])

d.h1('3.  Exclusions')
d.para('Confidential Information does not include information that:')
d.numbered([
    'Is or becomes publicly known through no breach of this Agreement by the Receiving Party;',
    'Was rightfully in the Receiving Party\'s possession before disclosure by the Disclosing Party;',
    'Is rightfully obtained by the Receiving Party from a third party without restriction;',
    'Is independently developed by the Receiving Party without use of the Confidential Information; or',
    'Is required to be disclosed by law, regulation, or court order — provided the Receiving Party gives the Disclosing Party prompt written notice and cooperates in seeking a protective order.',
])

d.pb()
d.h1('4.  Obligations of the Receiving Party')
d.para('Each Party, as Receiving Party, agrees to:')
d.numbered([
    'Hold the Disclosing Party\'s Confidential Information in strict confidence using at least the same degree of care it uses to protect its own confidential information, but in no event less than reasonable care;',
    'Not disclose Confidential Information to any third party without the prior written consent of the Disclosing Party;',
    'Use the Confidential Information solely for the Purpose described in this Agreement and for no other purpose;',
    'Limit access to Confidential Information to its employees, contractors, and advisors who have a need to know for the Purpose and who are bound by confidentiality obligations no less restrictive than those in this Agreement;',
    'Promptly notify the Disclosing Party in writing upon becoming aware of any actual or suspected unauthorised disclosure or use of Confidential Information.',
])

d.h1('5.  Term')
d.para(
    'This Agreement shall remain in effect for a period of three (3) years from the Effective Date. '
    'The obligations with respect to Confidential Information that constitutes a trade secret shall '
    'survive indefinitely until such information no longer qualifies as a trade secret under applicable law.'
)

d.h1('6.  Return or Destruction of Information')
d.para(
    'Upon written request by the Disclosing Party, or upon termination of this Agreement, '
    'the Receiving Party shall promptly return or destroy all copies of Confidential Information '
    'in its possession or control and, upon request, certify in writing that it has done so.'
)

d.h1('7.  No Licence or Rights')
d.para(
    'Nothing in this Agreement grants either Party any licence, right, title, or interest in or to '
    'any Confidential Information, intellectual property, or technology of the other Party, '
    'except the limited right to use such Confidential Information for the Purpose.'
)

d.h1('8.  No Obligation')
d.para(
    'Nothing in this Agreement obligates either Party to enter into any further agreement, '
    'proceed with any transaction, or disclose any particular information.'
)

d.h1('9.  Remedies')
d.para(
    'Each Party acknowledges that breach of this Agreement may cause irreparable harm '
    'for which monetary damages would be an inadequate remedy, '
    'and that the non-breaching Party shall be entitled to seek injunctive or other equitable relief '
    'in addition to all other remedies available at law or in equity.'
)

d.h1('10.  Governing Law and Jurisdiction')
d.para(
    f'This Agreement shall be governed by and construed in accordance with the laws of {COUNTRY}. '
    'Each Party irrevocably submits to the exclusive jurisdiction of the courts of England and Wales '
    'for any dispute arising out of or in connection with this Agreement.'
)

d.h1('11.  Entire Agreement')
d.para(
    'This Agreement constitutes the entire agreement between the Parties with respect to its subject matter '
    'and supersedes all prior agreements, understandings, and negotiations, whether written or oral.'
)

d.h1('12.  Signatures')
d.para('By signing below, each Party agrees to be bound by the terms of this Agreement.')
d.para('')
d.sig_block(OWNER, f'Owner, {TRADING_AS}', show_witness=True)
d.para('')
d.divider()
d.para('')
d.para('Party B:', bold=True)
d.sig_block('___________________________________________________________', 'Title / Role', show_witness=True)

d.footer(f'EchoPose Mutual NDA · {OWNER} · Governing Law: England & Wales')
d.save()


# ══════════════════════════════════════════════════════════════════════════════
#  ONE-WAY NDA  (EchoPose discloses to third party)
# ══════════════════════════════════════════════════════════════════════════════
print('Building One-Way NDA...')
d = Doc('EchoPose_OneWay_NDA.docx', BLUE)
d.cover(
    'One-Way Non-Disclosure Agreement',
    'For use when EchoPose shares confidential information with a third party',
    'LEGAL TEMPLATE  |  Version 1.0  |  Governing Law: England & Wales'
)

d.callout(
    'Use this agreement when sharing EchoPose technology details, business plans, or customer data '
    'with a contractor, employee, consultant, or third party who does not need to share their own confidential information back. '
    'Only the Receiving Party has obligations under this agreement.',
    BLUE, 'WHEN TO USE THIS DOCUMENT'
)

d.para('ONE-WAY NON-DISCLOSURE AGREEMENT', bold=True, size=14, color=NAVY)
d.para('')
d.para('This Non-Disclosure Agreement ("Agreement") is entered into as of the date last signed below ("Effective Date") between:', size=11)
d.para('')
d.para(f'Disclosing Party:  {OWNER}, trading as {TRADING_AS}', bold=True)
d.para(f'                   Contact: {EMAIL}')
d.para('')
d.para('Receiving Party:   ___________________________________________________________', bold=True)
d.para('                   Name:  ___________________________________________________')
d.para('                   Address:  ________________________________________________')
d.para('                   Contact:  ________________________________________________')
d.para('')

d.h1('1.  Background')
d.para(
    f'The Disclosing Party ({OWNER}, trading as {TRADING_AS}) intends to disclose certain '
    'confidential and proprietary information to the Receiving Party for the purpose of: '
    '____________________________________________________________________________ ("the Purpose"). '
    'The Receiving Party agrees to receive this information on the terms set out below.'
)

d.h1('2.  Confidential Information')
d.para(
    'For the purposes of this Agreement, "Confidential Information" means all information disclosed '
    f'by {TRADING_AS} to the Receiving Party, in any form or medium, including but not limited to:'
)
d.bullet([
    'Source code, algorithms, AI models, signal processing techniques, and technical architecture',
    'Business strategy, customer lists, pricing, financial data, and commercial plans',
    'Product roadmaps, unreleased features, and development plans',
    'Customer and user data, including any personal data processed by EchoPose systems',
    'Any information marked "Confidential", "Proprietary", or similar',
])

d.h1('3.  Obligations')
d.para('The Receiving Party agrees to:')
d.numbered([
    'Keep all Confidential Information strictly confidential and not disclose it to any third party;',
    'Use the Confidential Information only for the Purpose stated above;',
    'Protect the Confidential Information with the same care used to protect its own confidential information, and in no case less than reasonable care;',
    'Not copy, reproduce, reverse engineer, or attempt to derive the composition or underlying information of any Confidential Information;',
    'Not use Confidential Information to compete with EchoPose or to assist any third party in competing with EchoPose;',
    'Promptly notify EchoPose of any actual or suspected breach of confidentiality.',
])

d.h1('4.  Exclusions')
d.para('Obligations under this Agreement do not apply to information that the Receiving Party can demonstrate:')
d.numbered([
    'Was already known to the Receiving Party at the time of disclosure;',
    'Is or becomes publicly known through no act or omission of the Receiving Party;',
    'Is rightfully received from a third party without restriction on disclosure;',
    'Is required to be disclosed by law or court order (with prompt prior written notice to EchoPose).',
])

d.h1('5.  Intellectual Property')
d.para(
    'Nothing in this Agreement transfers any intellectual property rights to the Receiving Party. '
    'All Confidential Information remains the property of EchoPose '
    f'(Muhammed Shazin Sadhik Kunhi Parambath). '
    'The Receiving Party acquires no licence or rights by virtue of disclosure.'
)

d.h1('6.  Term')
d.para(
    'This Agreement is effective from the Effective Date and shall continue for three (3) years, '
    'unless terminated earlier by mutual written agreement. '
    'Obligations with respect to trade secrets shall survive indefinitely.'
)

d.h1('7.  Consequences of Breach')
d.para(
    'The Receiving Party acknowledges that any breach of this Agreement would cause '
    'irreparable damage to EchoPose for which monetary compensation would be insufficient. '
    f'{TRADING_AS} shall be entitled to seek injunctive relief, specific performance, '
    'and any other remedy available at law or in equity, without the requirement to post a bond.'
)

d.h1('8.  Governing Law')
d.para(
    f'This Agreement is governed by the laws of {COUNTRY}. '
    'Any disputes shall be subject to the exclusive jurisdiction of the courts of England and Wales.'
)

d.h1('9.  Signatures')
d.para('')
d.sig_block(OWNER, f'Owner, {TRADING_AS}', show_witness=False)
d.para('')
d.divider()
d.para('')
d.para('Receiving Party:', bold=True)
d.sig_block('___________________________________________________________', 'Title / Role', show_witness=True)

d.footer(f'EchoPose One-Way NDA · {OWNER} · Governing Law: England & Wales')
d.save()

print(f'\nAll HR & Legal documents saved to {OUT}/')
