"""
EchoPose — Legal Operations Documents Generator
Produces 5 documents in docs/legal/
Run: python scripts/generate_legal_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT    = 'docs/legal'
os.makedirs(OUT, exist_ok=True)

OWNER      = 'Muhammed Shazin Sadhik Kunhi Parambath'
TRADING_AS = 'EchoPose'
EMAIL      = 'shazin2889@gmail.com'
WEBSITE    = 'github.com/shaz-in-dev/EchoPose'  # update to echopose.io when domain is live
ADDRESS    = '[Your registered address]' # update when known
COUNTRY    = 'England and Wales'
DATE       = 'April 2026'

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
BLUE   = RGBColor(0x1B, 0x6C, 0xA8)
TEAL   = RGBColor(0x00, 0x87, 0x8A)
GREEN  = RGBColor(0x1A, 0x73, 0x48)
RED    = RGBColor(0xC0, 0x20, 0x20)
ORANGE = RGBColor(0xD9, 0x6B, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF2, 0xF4, 0xF7)
MGREY  = RGBColor(0x99, 0xA3, 0xB0)
DKGREY = RGBColor(0x3A, 0x3F, 0x4A)

def hx(c): return f'{c[0]:02X}{c[1]:02X}{c[2]:02X}'

# ── Doc helper ────────────────────────────────────────────────────────────────
class Doc:
    def __init__(self, filename, accent=BLUE):
        self.d = Document(); self.fn = filename; self.ac = accent
        s = self.d.sections[0]
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin  = s.bottom_margin = Cm(2.5)

    def save(self):
        p = os.path.join(OUT, self.fn)
        self.d.save(p); print(f'  Saved: {p}')

    def _shd(self, cell, rgb):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        e = OxmlElement('w:shd')
        e.set(qn('w:val'), 'clear'); e.set(qn('w:color'), 'auto')
        e.set(qn('w:fill'), hx(rgb)); pr.append(e)

    def pb(self): self.d.add_page_break()

    def cover(self, title, subtitle, doc_label):
        t = self.d.add_table(1, 1)
        c = t.cell(0, 0); self._shd(c, NAVY); c.width = Inches(6)
        for txt, sz, bold, col in [
            (TRADING_AS, 28, True, TEAL),
            (title, 18, True, WHITE),
            ('', 5, False, WHITE),
            (subtitle, 11, False, RGBColor(0xB0, 0xD0, 0xFF)),
            ('', 5, False, WHITE),
            (doc_label, 9, True, ORANGE),
            (f'{OWNER}  ·  {EMAIL}', 8, False, MGREY),
            (f'Version 1.0  ·  {DATE}  ·  Governing Law: {COUNTRY}', 8, False, MGREY),
        ]:
            p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt); r.font.size = Pt(sz)
            r.font.bold = bold; r.font.color.rgb = col
        self.d.add_paragraph(); self.pb()

    def banner(self, text, color=None):
        color = color or self.ac
        t = self.d.add_table(1, 1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
        c = t.cell(0, 0); self._shd(c, color); c.width = Inches(5.8)
        p = c.paragraphs[0]
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text); r.font.bold = True
        r.font.size = Pt(12); r.font.color.rgb = WHITE
        self.d.add_paragraph()

    def h1(self, txt, color=None):
        p = self.d.add_heading(txt, 1)
        for r in p.runs:
            r.font.color.rgb = color or self.ac; r.font.bold = True

    def h2(self, txt, color=None):
        p = self.d.add_heading(txt, 2)
        for r in p.runs:
            r.font.color.rgb = color or self.ac; r.font.bold = True

    def h3(self, txt):
        p = self.d.add_heading(txt, 3)
        for r in p.runs: r.font.color.rgb = NAVY

    def para(self, txt='', bold=False, italic=False, color=None, size=11):
        p = self.d.add_paragraph()
        if txt:
            r = p.add_run(txt); r.font.size = Pt(size)
            r.font.bold = bold; r.font.italic = italic
            if color: r.font.color.rgb = color
        return p

    def bullet(self, items, bold_prefix=None):
        for item in items:
            p = self.d.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            if bold_prefix and item.startswith(bold_prefix):
                parts = item.split(':', 1)
                if len(parts) == 2:
                    rb = p.add_run(parts[0] + ':'); rb.font.bold = True; rb.font.size = Pt(10.5)
                    rn = p.add_run(parts[1]); rn.font.size = Pt(10.5)
                    continue
            r = p.add_run(item); r.font.size = Pt(10.5)

    def numbered(self, items):
        for item in items:
            p = self.d.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(item); r.font.size = Pt(10.5)

    def sub_numbered(self, items, prefix='a'):
        letters = 'abcdefghijklmnopqrstuvwxyz'
        for i, item in enumerate(items):
            p = self.d.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.8)
            r = p.add_run(f'({letters[i]})  {item}'); r.font.size = Pt(10.5)

    def table(self, headers, rows, hcol=None):
        hcol = hcol or self.ac
        t = self.d.add_table(1 + len(rows), len(headers))
        t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; self._shd(c, hcol)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h); r.font.bold = True
            r.font.size = Pt(10); r.font.color.rgb = WHITE
        for ri, row in enumerate(rows):
            bg = LGREY if ri % 2 == 0 else WHITE
            for ci, val in enumerate(row):
                c = t.rows[ri + 1].cells[ci]; self._shd(c, bg)
                p = c.paragraphs[0]; r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.d.add_paragraph()

    def callout(self, txt, color=None, label=''):
        color = color or self.ac
        t = self.d.add_table(1, 1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
        c = t.cell(0, 0); self._shd(c, color); c.width = Inches(5.8)
        p = c.paragraphs[0]
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(5)
        if label:
            rl = p.add_run(label + '  '); rl.font.bold = True
            rl.font.size = Pt(10); rl.font.color.rgb = WHITE
        r = p.add_run(txt); r.font.size = Pt(10); r.font.color.rgb = WHITE
        self.d.add_paragraph()

    def sig(self, name, role, company='', witness=False):
        self.para('')
        lines = [
            f'Signed:      _______________________________________________',
            f'',
            f'Full Name:   {name}',
            f'',
            f'Role:        {role}',
        ]
        if company:
            lines.append(f'Company:     {company}')
        lines += ['', 'Date:        _______________________________________________', '']
        if witness:
            lines += [
                'Witness Signature:  ________________________________________',
                'Witness Full Name:  ________________________________________',
                'Witness Address:    ________________________________________',
                '',
            ]
        for line in lines:
            p = self.d.add_paragraph()
            r = p.add_run(line); r.font.size = Pt(10.5)
            if 'Signed:' in line or 'Full Name:' in line or 'Role:' in line:
                r.font.bold = True

    def divider(self):
        p = self.d.add_paragraph()
        r = p.add_run('─' * 85); r.font.size = Pt(7); r.font.color.rgb = MGREY

    def footer(self, text):
        self.divider()
        p = self.d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(8); r.font.color.rgb = MGREY


# ══════════════════════════════════════════════════════════════════════════════
#  1.  TERMS OF SERVICE
# ══════════════════════════════════════════════════════════════════════════════
print('Building Terms of Service...')
d = Doc('EchoPose_Terms_of_Service.docx', NAVY)
d.cover(
    'Terms of Service',
    'The legal agreement between EchoPose and every customer who uses the platform',
    'LEGAL DOCUMENT  |  Effective: April 2026'
)

d.callout(
    'Please read these Terms carefully before using EchoPose. '
    'By purchasing a subscription, installing the hardware, or accessing the dashboard, '
    'you agree to be bound by these Terms. If you do not agree, do not use the service.',
    RED, '⚠ IMPORTANT'
)

d.para(
    f'These Terms of Service ("Terms") constitute a legally binding agreement between you ("Customer", "you") '
    f'and {OWNER}, trading as {TRADING_AS} ("EchoPose", "we", "us", "our"), '
    f'a sole trader operating under the laws of {COUNTRY}. '
    f'Our contact email is {EMAIL}.',
    size=10
)

# ─── Section 1 ────────────────────────────────────────────────────────────────
d.h1('1.  The Service')
d.h2('1.1  What EchoPose Provides')
d.para(
    'EchoPose provides a WiFi Channel State Information (CSI) sensing platform that uses '
    'radio frequency signals to detect human presence, movement, posture, and certain physiological indicators '
    'within a monitored space ("the Service"). '
    'The Service comprises hardware sensors, firmware, a server-side AI inference engine, '
    'a web-based dashboard, and alert notifications.'
)

d.h2('1.2  What EchoPose is NOT')
d.callout(
    'EchoPose is NOT a medical device. It is NOT regulated under the UK Medical Devices Regulations 2002. '
    'It is NOT a substitute for professional medical advice, diagnosis, or treatment. '
    'It is NOT a substitute for emergency services. In an emergency, always call 999.',
    RED, '⚠ CRITICAL DISCLAIMER'
)
d.para(
    'EchoPose is a home safety and activity monitoring tool only. '
    'Vital sign estimates (breathing rate, approximate heart rate) are indicative only '
    'and must not be relied upon for any clinical or medical purpose. '
    'Fall detection alerts are best-effort and are not guaranteed to be triggered in every incident.'
)

d.h2('1.3  Service Availability')
d.para(
    'We target 99% uptime for cloud-hosted services but do not guarantee uninterrupted availability. '
    'Scheduled maintenance will be communicated with at least 24 hours notice where possible. '
    'On-premise deployments are not subject to uptime guarantees from EchoPose; '
    'availability is determined by the customer\'s own infrastructure.'
)

# ─── Section 2 ────────────────────────────────────────────────────────────────
d.h1('2.  Accounts and Eligibility')
d.h2('2.1  Eligibility')
d.para(
    'You must be at least 18 years of age and capable of entering into a binding contract '
    'under the laws of England and Wales to use the Service. '
    'By agreeing to these Terms, you represent and warrant that you meet this requirement.'
)

d.h2('2.2  Account Registration')
d.para(
    'You are responsible for maintaining the confidentiality of your account credentials. '
    'You must notify us immediately at ' + EMAIL + ' if you suspect unauthorised access to your account. '
    'We are not liable for any loss or damage arising from your failure to keep credentials secure.'
)

d.h2('2.3  Acceptable Use')
d.para('You agree that you will NOT use the Service to:')
d.numbered([
    'Monitor any individual without their explicit informed consent (or, where they lack capacity, the consent of their legal representative);',
    'Violate any applicable law, including data protection laws, surveillance laws, or privacy laws;',
    'Monitor public spaces, common areas, or any space where the monitored individuals have a reasonable expectation not to be monitored;',
    'Conduct surveillance for purposes other than the safety and wellbeing of the consented individual;',
    'Reverse engineer, decompile, or attempt to extract the source code of any EchoPose software;',
    'Resell, sublicense, or otherwise make the Service available to third parties except as permitted under a separate Reseller Agreement;',
    'Interfere with or disrupt the integrity or performance of the Service or its infrastructure.',
])

# ─── Section 3 ────────────────────────────────────────────────────────────────
d.pb()
d.h1('3.  Subscriptions, Payment and Cancellation')
d.h2('3.1  Subscription Tiers')
d.table(
    ['Tier', 'Price', 'Billing', 'Rooms Covered', 'Notice to Cancel'],
    [
        ['EchoPose Home',  '£49/month',  'Monthly, in advance', '1', '30 days written notice'],
        ['EchoPose Care',  '£199/month', 'Monthly, in advance', 'Up to 5', '30 days written notice'],
        ['EchoPose Pro',   '£499/month', 'Monthly, in advance', 'Up to 20', '30 days written notice'],
        ['On-Premise License', 'As quoted', 'One-time + annual support', 'As per agreement', 'N/A — perpetual license'],
    ]
)

d.h2('3.2  Payment')
d.para(
    'Subscriptions are billed monthly in advance via Stripe or as otherwise agreed. '
    'All prices are in GBP (£) and exclude VAT. EchoPose is not currently VAT registered. '
    'Should EchoPose become VAT registered, prices will be adjusted accordingly with 30 days notice. '
    'Payment must be made within 14 days of invoice. '
    'Overdue accounts will result in suspension of access until payment is received.'
)

d.h2('3.3  Cancellation and Refunds')
d.para(
    'You may cancel your subscription at any time by giving 30 days written notice to ' + EMAIL + '. '
    'Cancellation takes effect at the end of the current billing period. '
    'No refunds are given for partial months. '
    'Hardware purchases are non-refundable once the hardware has been activated, '
    'except where required by the Consumer Rights Act 2015 (for consumers) or as otherwise agreed in writing.'
)

d.h2('3.4  Price Changes')
d.para(
    'We reserve the right to change subscription prices. '
    'We will give you at least 30 days written notice of any price increase. '
    'If you do not wish to accept the new price, you may cancel before the new price takes effect.'
)

# ─── Section 4 ────────────────────────────────────────────────────────────────
d.h1('4.  Data Protection and Privacy')
d.para(
    'Our collection and use of personal data is governed by our Privacy Policy, '
    'which is incorporated into these Terms by reference. '
    'Business customers (care agencies, care homes, enterprises) must enter into a '
    'separate Data Processing Agreement (DPA) with EchoPose before the Service processes '
    'personal data belonging to their clients or employees. '
    'You are responsible for ensuring that all individuals monitored using the Service '
    'have given their informed consent in accordance with applicable law.'
)

# ─── Section 5 ────────────────────────────────────────────────────────────────
d.h1('5.  Intellectual Property')
d.para(
    'EchoPose and all associated technology, software, algorithms, models, firmware, '
    'documentation, trademarks, and content ("IP") are the exclusive property of '
    f'{OWNER}. '
    'These Terms grant you a limited, non-exclusive, non-transferable licence to use '
    'the Service solely for your own internal purposes in accordance with these Terms. '
    'No other rights are granted. '
    'You may not copy, modify, distribute, or create derivative works from any EchoPose IP '
    'without prior written consent.'
)

# ─── Section 6 ────────────────────────────────────────────────────────────────
d.pb()
d.h1('6.  Limitation of Liability')
d.callout(
    'This section significantly limits our liability to you. Please read it carefully.',
    NAVY, '⚠ IMPORTANT'
)

d.h2('6.1  Exclusion of Consequential Loss')
d.para(
    'To the fullest extent permitted by law, EchoPose shall not be liable for any '
    'indirect, incidental, special, consequential, or punitive damages, '
    'including but not limited to: loss of profit, loss of data, loss of business, '
    'personal injury, or property damage arising from or in connection with the use of the Service, '
    'even if EchoPose has been advised of the possibility of such damages.'
)

d.h2('6.2  Cap on Liability')
d.para(
    'EchoPose\'s total aggregate liability to you in connection with these Terms '
    'shall not exceed the total amount paid by you to EchoPose '
    'in the three (3) months immediately preceding the claim.'
)

d.h2('6.3  Fall Detection and Vital Signs — Specific Disclaimer')
d.para(
    'EchoPose\'s fall detection and vital sign monitoring features are provided on a best-effort basis. '
    'They are not guaranteed to detect every fall or health event. '
    'EchoPose expressly disclaims all liability for any personal injury, death, or harm '
    'arising from a failure to detect, or a delayed detection of, any fall or health event. '
    'You acknowledge that EchoPose is not a medical device and must not be relied upon as a primary safety mechanism.'
)

d.h2('6.4  Consumer Rights')
d.para(
    'Nothing in these Terms excludes or limits EchoPose\'s liability for: '
    'death or personal injury caused by our negligence; fraud or fraudulent misrepresentation; '
    'or any other liability that cannot be excluded or limited under applicable law '
    '(including the Consumer Rights Act 2015 where applicable).'
)

# ─── Section 7 ────────────────────────────────────────────────────────────────
d.h1('7.  Hardware')
d.para(
    'Hardware sold by EchoPose is warranted against manufacturing defects for 12 months from purchase. '
    'This warranty does not cover damage caused by misuse, unauthorised modification, or '
    'failure to follow installation instructions. '
    'You are responsible for ensuring hardware is installed in a safe and appropriate manner. '
    'EchoPose does not accept liability for damage to property caused by improper installation.'
)

# ─── Section 8 ────────────────────────────────────────────────────────────────
d.h1('8.  Termination')
d.para(
    'Either party may terminate these Terms by cancelling the subscription as described in Section 3.3. '
    'EchoPose may suspend or terminate your access immediately and without notice if: '
    'you materially breach these Terms; you fail to pay amounts due; '
    'we are required to do so by law; or we have reasonable grounds to believe your use '
    'is unlawful or harmful to others. '
    'On termination, your licence to use the Service ends immediately. '
    'Sections 5, 6, 9, and 10 survive termination.'
)

# ─── Section 9 ────────────────────────────────────────────────────────────────
d.h1('9.  Governing Law and Disputes')
d.para(
    f'These Terms are governed by the laws of {COUNTRY}. '
    'Both parties agree to attempt to resolve any dispute informally in the first instance '
    'by contacting EchoPose at ' + EMAIL + '. '
    'If the dispute cannot be resolved within 30 days, '
    'it shall be subject to the exclusive jurisdiction of the courts of England and Wales. '
    'Nothing prevents either party from seeking emergency injunctive relief in any competent court.'
)

# ─── Section 10 ────────────────────────────────────────────────────────────────
d.h1('10.  General')
d.table(
    ['Clause', 'Meaning'],
    [
        ['Entire Agreement', 'These Terms (plus Privacy Policy and any DPA) constitute the entire agreement between the parties and supersede all prior understandings.'],
        ['Severability', 'If any provision is found to be unenforceable, the remaining provisions continue in full force.'],
        ['Waiver', 'Failure to enforce any provision is not a waiver of the right to enforce it in the future.'],
        ['Assignment', 'You may not assign your rights under these Terms without our written consent. We may assign our rights on notice to you.'],
        ['Notices', f'Written notices to EchoPose must be sent to {EMAIL}. We may give notices to you via the email address on your account.'],
        ['Updates to Terms', 'We may update these Terms with 30 days notice. Continued use after that date constitutes acceptance.'],
        ['Force Majeure', 'EchoPose is not liable for failures caused by events outside our reasonable control (natural disasters, network outages, government action).'],
    ]
)

d.callout(
    'These Terms should be reviewed by a qualified solicitor before being relied upon commercially. '
    'They have been drafted as a starting framework under English law.',
    ORANGE, '⚠ LEGAL NOTE'
)
d.footer(f'EchoPose Terms of Service · {OWNER} · {DATE} · Governing Law: {COUNTRY}')
d.save()


# ══════════════════════════════════════════════════════════════════════════════
#  2.  PRIVACY POLICY
# ══════════════════════════════════════════════════════════════════════════════
print('Building Privacy Policy...')
d = Doc('EchoPose_Privacy_Policy.docx', TEAL)
d.cover(
    'Privacy Policy',
    'How EchoPose collects, uses, stores, and protects your personal data',
    'GDPR COMPLIANT  |  UK Data Protection Act 2018  |  Effective: April 2026'
)

d.para(
    f'This Privacy Policy explains how {OWNER}, trading as {TRADING_AS} '
    f'("EchoPose", "we", "us", "our"), collects and uses personal data in connection '
    f'with the EchoPose platform and services. '
    f'EchoPose is the data controller for personal data processed under this policy. '
    f'Contact: {EMAIL}.',
    size=10
)

d.h1('1.  Who We Are')
d.para(
    f'Data Controller: {OWNER}, trading as {TRADING_AS}. '
    f'Contact: {EMAIL}. '
    f'EchoPose is registered with the Information Commissioner\'s Office (ICO) as a data controller. '
    f'ICO registration number: [Insert once registered at ico.org.uk].'
)

d.h1('2.  What Personal Data We Collect')
d.table(
    ['Category', 'Specific Data', 'Source', 'Is it Special Category?'],
    [
        ['Account Data',
         'Name, email address, billing address, payment method token',
         'Provided by the account holder at registration',
         'No'],
        ['Monitored Person Data — Movement',
         '3D skeleton keypoint coordinates (17 body points: x, y, z, confidence), activity classification (walking / sitting / lying / etc.), room occupancy state',
         'Inferred from WiFi CSI signals by EchoPose AI engine',
         'Potentially — may constitute health-related data'],
        ['Monitored Person Data — Physiological',
         'Estimated breathing rate (breaths per minute), approximate heart rate (beats per minute), sleep stage classification, stress/emotion score',
         'Inferred from WiFi CSI micro-Doppler signal analysis',
         'YES — Special Category Health Data (Article 9 UK GDPR). Requires explicit consent.'],
        ['Alert and Event Logs',
         'Fall event timestamp, location ID, severity score; vital anomaly events; alert recipient and delivery status',
         'Generated automatically when events are detected',
         'Yes — linked to health state of monitored person'],
        ['Device and Technical Data',
         'IP address, browser type, dashboard session data, hardware node IDs, firmware version',
         'Collected automatically when using the dashboard',
         'No'],
        ['Communication Data',
         'Emails sent to/from EchoPose support, content of support tickets',
         'Provided by the user directly',
         'No'],
    ]
)

d.pb()
d.h1('3.  Why We Process Your Data (Lawful Basis)')
d.table(
    ['Processing Activity', 'Lawful Basis (UK GDPR)', 'Detail'],
    [
        ['Providing the subscription service and billing',
         'Article 6(1)(b) — Performance of a contract',
         'Necessary to deliver the service you have paid for'],
        ['Monitoring movement and activity of consented individuals',
         'Article 6(1)(f) — Legitimate interests\n(safety and wellbeing)',
         'Legitimate interest assessment on file. Balanced against individual rights.'],
        ['Processing physiological/health data (breathing, heart rate, sleep)',
         'Article 9(2)(a) — Explicit consent',
         'Requires a signed consent form from the monitored person or their legal representative. This data will NOT be processed without this consent.'],
        ['Sending fall detection and health alerts',
         'Article 6(1)(f) — Legitimate interests (safety)',
         'Vital safety purpose. Balancing test on file.'],
        ['Fraud prevention and security',
         'Article 6(1)(f) — Legitimate interests',
         'Protecting the integrity of our systems'],
        ['Legal compliance (e.g. HMRC, ICO)',
         'Article 6(1)(c) — Legal obligation',
         'We retain billing records for 7 years as required by HMRC'],
        ['Marketing to existing customers',
         'Article 6(1)(f) — Legitimate interests',
         'Existing customer relationship. You may opt out at any time.'],
    ]
)

d.h1('4.  How Long We Keep Your Data')
d.table(
    ['Data Type', 'Retention Period', 'Reason'],
    [
        ['Raw WiFi CSI signals', '24 hours', 'Processed immediately; raw data deleted'],
        ['Skeleton / pose data', '90 days (configurable by account holder)', 'Allows review of recent events'],
        ['Fall and alert events', 'Duration of contract + 2 years', 'Audit trail and incident review'],
        ['Physiological estimates', '90 days (configurable)', 'Same as pose data'],
        ['Account and billing data', 'Duration of contract + 7 years', 'HMRC legal requirement'],
        ['Support communications', '3 years from last contact', 'Dispute resolution'],
        ['Device and technical logs', '90 days', 'Security and debugging'],
    ]
)

d.h1('5.  Who We Share Your Data With')
d.para('We do not sell your personal data. We share it only in the following circumstances:')
d.table(
    ['Recipient', 'Purpose', 'Safeguard'],
    [
        ['Stripe, Inc.', 'Payment processing', 'Standard Contractual Clauses (SCC). Stripe is PCI-DSS certified.'],
        ['SendGrid (Twilio)', 'Email alert delivery', 'DPA in place. EU/UK data centres available.'],
        ['Twilio, Inc.', 'SMS alert delivery', 'DPA in place.'],
        ['Cloud hosting provider\n(DigitalOcean / Hetzner)', 'Infrastructure hosting', 'DPA in place. UK/EU data centres used.'],
        ['Law enforcement / regulators', 'Where required by law or court order', 'Minimum disclosure; legal advice sought'],
        ['Business successor', 'In event of sale of the business', 'Customers will be notified in advance'],
    ]
)

d.pb()
d.h1('6.  Your Rights Under UK GDPR')
d.table(
    ['Right', 'What It Means', 'How to Exercise'],
    [
        ['Right to Access', 'Request a copy of all personal data we hold about you', f'Email {EMAIL} — we will respond within 30 days'],
        ['Right to Rectification', 'Ask us to correct inaccurate data', f'Email {EMAIL}'],
        ['Right to Erasure ("Right to be Forgotten")', 'Ask us to delete your data (subject to legal retention obligations)', f'Email {EMAIL}'],
        ['Right to Restrict Processing', 'Ask us to pause processing while a dispute is resolved', f'Email {EMAIL}'],
        ['Right to Data Portability', 'Receive your data in a machine-readable format', f'Email {EMAIL} — we provide JSON export'],
        ['Right to Object', 'Object to processing based on legitimate interests or for direct marketing', f'Email {EMAIL}. Marketing opt-out: unsubscribe link in every email'],
        ['Right to Withdraw Consent', 'Withdraw consent for health data processing at any time', 'This will stop health data collection immediately. Contact {EMAIL}'],
        ['Right to Complain', 'Lodge a complaint with the ICO', 'ico.org.uk — Tel: 0303 123 1113'],
    ]
)

d.h1('7.  Data Security')
d.para(
    'We implement appropriate technical and organisational measures to protect personal data, including:'
)
d.bullet([
    'All data in transit is encrypted using TLS 1.3',
    'All data at rest is encrypted using AES-256 (Fernet)',
    'Access to personal data is restricted to authorised personnel only',
    'API access is protected by token authentication and rate limiting',
    'No raw images or video are ever stored by EchoPose',
    'Regular security reviews are conducted',
    'In the event of a data breach, we will notify the ICO within 72 hours and affected individuals without undue delay',
])

d.h1('8.  International Transfers')
d.para(
    'Where personal data is transferred outside the UK/EEA (e.g. to Stripe or Twilio in the US), '
    'we ensure appropriate safeguards are in place, '
    'such as UK International Data Transfer Agreements (IDTAs) or Standard Contractual Clauses (SCCs). '
    'You may request a copy of the relevant safeguards by emailing ' + EMAIL + '.'
)

d.h1('9.  Cookies')
d.para(
    'Our dashboard uses essential session cookies to maintain your login state. '
    'We do not use third-party tracking cookies or advertising cookies. '
    'You may disable cookies in your browser settings, but this may affect dashboard functionality.'
)

d.h1('10.  Changes to this Policy')
d.para(
    'We may update this Privacy Policy from time to time. '
    'We will notify you of significant changes by email at least 14 days before they take effect. '
    'Continued use of the Service after that date constitutes acceptance of the updated policy. '
    'The current version is always available at ' + WEBSITE + '/privacy.'
)

d.h1('11.  Contact')
d.para(f'Data Controller: {OWNER}, trading as {TRADING_AS}')
d.para(f'Email: {EMAIL}')
d.para(f'Website: {WEBSITE}')
d.para(f'ICO Registration Number: [Insert on registration]')
d.footer(f'EchoPose Privacy Policy · {OWNER} · {DATE} · UK GDPR Compliant')
d.save()


# ══════════════════════════════════════════════════════════════════════════════
#  3.  DATA PROCESSING AGREEMENT (DPA)
# ══════════════════════════════════════════════════════════════════════════════
print('Building Data Processing Agreement...')
d = Doc('EchoPose_Data_Processing_Agreement.docx', NAVY)
d.cover(
    'Data Processing Agreement',
    'Required when EchoPose processes personal data on behalf of a business customer',
    'UK GDPR Article 28  |  LEGAL TEMPLATE  |  Version 1.0'
)

d.callout(
    'This DPA is required when a business customer (care agency, care home, enterprise) '
    'uses EchoPose to process personal data belonging to their clients, residents, or employees. '
    'The business customer is the Controller; EchoPose is the Processor.',
    BLUE, 'WHEN TO USE THIS DOCUMENT'
)

d.para('DATA PROCESSING AGREEMENT', bold=True, size=14, color=NAVY)
d.para('')
d.para('This Data Processing Agreement ("DPA") is entered into between:', size=11)
d.para('')
d.para(f'Data Controller ("Controller"):  ___________________________________________', bold=True)
d.para(f'    Company / Organisation Name:  _____________________________________________')
d.para(f'    Registered Address:  ______________________________________________________')
d.para(f'    Data Protection Contact:  _________________________________________________')
d.para(f'    ICO Registration Number:  _________________________________________________')
d.para('')
d.para(f'Data Processor ("Processor"):  {OWNER}, trading as {TRADING_AS}', bold=True)
d.para(f'    Contact: {EMAIL}')
d.para('')
d.para('Together referred to as "the Parties".')

d.h1('1.  Definitions')
d.table(
    ['Term', 'Meaning'],
    [
        ['"Processing"', 'Any operation on personal data including collection, storage, analysis, transmission, and deletion'],
        ['"Personal Data"', 'Any information relating to an identified or identifiable natural person as defined in UK GDPR Article 4'],
        ['"Special Category Data"', 'Data revealing health information, biometric data, or other categories listed in UK GDPR Article 9'],
        ['"Data Subject"', 'The individual whose personal data is being processed (e.g. a care home resident being monitored)'],
        ['"Sub-processor"', 'A third party engaged by EchoPose to assist in processing (e.g. cloud hosting provider)'],
        ['"UK GDPR"', 'The UK General Data Protection Regulation as retained in UK law by the European Union (Withdrawal) Act 2018'],
    ]
)

d.h1('2.  Subject Matter, Duration, and Nature of Processing')
d.table(
    ['Element', 'Detail'],
    [
        ['Subject Matter', 'Processing of personal data in connection with the EchoPose human sensing service'],
        ['Duration', 'For the term of the service agreement between Controller and Processor, plus any retention period required by law'],
        ['Nature of Processing', 'Collection, analysis, storage, and transmission of WiFi CSI-derived personal data; generation of alerts and reports'],
        ['Purpose', 'Safety monitoring — fall detection, activity monitoring, vital sign estimation — as instructed by the Controller'],
        ['Types of Personal Data', 'Movement data (skeleton/pose), physiological estimates (breathing rate, heart rate), activity logs, fall/event alerts'],
        ['Special Category Data', 'Health-related data (physiological estimates, fall events). Processed only with explicit consent of data subjects.'],
        ['Categories of Data Subjects', 'Individuals (residents, clients, employees) within spaces monitored by the Controller using EchoPose hardware'],
    ]
)

d.pb()
d.h1('3.  Obligations of the Processor (EchoPose)')
d.para('EchoPose agrees to:')
d.numbered([
    'Process personal data only on documented instructions from the Controller, including with regard to international transfers, unless required to do so by law;',
    'Ensure that persons authorised to process the personal data are committed to confidentiality;',
    'Implement and maintain appropriate technical and organisational measures to ensure a level of security appropriate to the risk (Article 32 UK GDPR);',
    'Respect the conditions for engaging sub-processors set out in Clause 4;',
    'Assist the Controller, by appropriate technical and organisational measures, in fulfilling obligations to respond to data subject rights requests;',
    'Assist the Controller in ensuring compliance with Articles 32–36 (security, breach notification, DPIA) of UK GDPR;',
    'Delete or return all personal data to the Controller at the end of the service, and delete existing copies unless UK law requires storage;',
    'Make available to the Controller all information necessary to demonstrate compliance, and allow for audits at reasonable notice.',
])

d.h1('4.  Sub-processors')
d.para('EchoPose currently uses the following sub-processors:')
d.table(
    ['Sub-processor', 'Location', 'Purpose', 'Safeguard'],
    [
        ['DigitalOcean / Hetzner', 'UK / EU', 'Infrastructure hosting and data storage', 'DPA in place, EU/UK data centres'],
        ['Stripe, Inc.', 'US', 'Payment processing (account data only, not monitoring data)', 'SCC / IDTA in place'],
        ['Twilio SendGrid', 'US / EU', 'Email alert delivery', 'DPA in place'],
        ['Twilio, Inc.', 'US', 'SMS alert delivery', 'DPA in place'],
    ]
)
d.para(
    'EchoPose will notify the Controller of any intended changes to sub-processors '
    'by giving at least 14 days written notice. '
    'The Controller may object to new sub-processors within 14 days of notice. '
    'If no objection is received, the new sub-processor is deemed accepted.'
)

d.h1('5.  Data Security Measures')
d.para('EchoPose implements the following technical and organisational security measures (Article 32):')
d.bullet([
    'Encryption in transit: TLS 1.3 for all data transmissions',
    'Encryption at rest: AES-256 (Fernet) for all stored personal data',
    'Access control: Role-based access, API token authentication, rate limiting',
    'No raw images or video stored at any time',
    'Regular patching and security updates',
    'Incident response plan in place; ICO notification within 72 hours of discovering a breach',
    'Pseudonymisation of monitoring data where technically feasible',
])

d.h1('6.  Data Breach Notification')
d.para(
    'EchoPose will notify the Controller without undue delay, and in any case within 48 hours, '
    'after becoming aware of a personal data breach involving data processed under this DPA. '
    'Notification will include: nature of the breach, categories and approximate number of data subjects affected, '
    'likely consequences, and measures taken or proposed to address the breach.'
)

d.h1('7.  Data Subject Rights')
d.para(
    'Where EchoPose receives a data subject rights request directly from a data subject '
    'whose data is processed under this DPA, EchoPose will promptly forward the request to the Controller '
    'and will not respond to the data subject directly without Controller authorisation. '
    'EchoPose will assist the Controller in responding to such requests within the applicable timescales.'
)

d.h1('8.  Deletion and Return of Data')
d.para(
    'On termination of the service agreement, EchoPose will, at the Controller\'s election, '
    'either return all personal data in JSON format or securely delete it. '
    'Deletion will be completed within 30 days of termination. '
    'EchoPose will provide written certification of deletion on request. '
    'This does not apply to data that EchoPose is legally required to retain.'
)

d.h1('9.  Audit Rights')
d.para(
    'EchoPose will, on reasonable notice (minimum 14 days), '
    'make available all information necessary to demonstrate compliance with this DPA '
    'and allow for audits by the Controller or a third party auditor appointed by the Controller. '
    'Audits will be conducted at the Controller\'s expense and in a manner that minimises disruption to operations.'
)

d.h1('10.  Governing Law')
d.para(
    f'This DPA is governed by the laws of {COUNTRY} '
    'and shall be construed in accordance with UK GDPR as applicable in the United Kingdom.'
)

d.h1('11.  Signatures')
d.para('By signing, both parties agree to the terms of this Data Processing Agreement.')
d.para('')
d.para('CONTROLLER:', bold=True)
d.sig('___________________________________________', 'Authorised Signatory', company='Controller organisation', witness=False)
d.divider()
d.para('PROCESSOR (EchoPose):', bold=True)
d.sig(OWNER, f'Owner, {TRADING_AS}', witness=False)
d.footer(f'EchoPose DPA · {OWNER} · {DATE} · UK GDPR Article 28')
d.save()


# ══════════════════════════════════════════════════════════════════════════════
#  4.  CONSENT FORM
# ══════════════════════════════════════════════════════════════════════════════
print('Building Consent Form...')
d = Doc('EchoPose_Consent_Form.docx', TEAL)
d.cover(
    'Monitoring Consent Form',
    'To be completed and signed before any individual is monitored by EchoPose',
    'UK GDPR Article 9 — Explicit Consent for Special Category Health Data'
)

d.callout(
    'This form MUST be completed and signed before EchoPose is activated in any space where '
    'an individual will be monitored. Processing health-related data (breathing rate, heart rate, '
    'activity) without this consent is unlawful under UK GDPR Article 9. '
    'Keep a signed copy on file for the duration of monitoring plus 2 years.',
    RED, '⚠ MANDATORY — DO NOT SKIP'
)

d.h1('PART A — About This Form')
d.para(
    'EchoPose is a WiFi-based safety monitoring system. '
    'It uses radio signals — not cameras — to monitor movement, posture, and certain physical indicators '
    'of the person in the monitored room. '
    'This form explains what data is collected, why, who can see it, and your rights. '
    'Please read it carefully before signing.'
)

d.h1('PART B — What Will Be Monitored')
d.table(
    ['What EchoPose Monitors', 'What This Means In Plain English', 'Is This Health Data?'],
    [
        ['Movement and position', 'Where you are in the room, whether you are standing, sitting, walking, or lying down', 'Potentially'],
        ['Breathing rate', 'How many breaths per minute — estimated from tiny chest movements', 'Yes — health-related'],
        ['Approximate heart rate', 'Rough estimate of heartbeats per minute from body micro-movement', 'Yes — health-related'],
        ['Fall detection', 'If a sudden fall occurs, the system will detect it and send an alert', 'Yes — health-related event'],
        ['Sleep patterns', 'Whether you are asleep and approximate sleep quality during night hours', 'Yes — health-related'],
        ['Activity levels', 'How active you have been throughout the day (daily summary report)', 'Potentially'],
    ]
)

d.callout(
    'EchoPose does NOT use cameras. No images or video are ever recorded. '
    'The system cannot identify you by your face. '
    'It only senses movement and position through WiFi signals.',
    GREEN, '✔ YOUR PRIVACY'
)

d.h1('PART C — Who Will See This Data')
d.para('Please tick all that apply and name the specific people or organisations:')
d.para('')
d.para('☐  Family member(s):  ___________________________________________________')
d.para('☐  Named carer(s):  _____________________________________________________')
d.para('☐  Care agency:  ________________________________________________________')
d.para('☐  Care home staff (named roles only):  ___________________________________')
d.para('☐  GP / healthcare professional:  ________________________________________')
d.para('☐  Other (specify):  _____________________________________________________')
d.para('')
d.callout(
    'Data will ONLY be shared with the people named above. '
    'It will not be sold, shared with advertisers, or disclosed to any other party '
    'without your separate written consent.',
    TEAL, '✔ YOUR DATA'
)

d.pb()
d.h1('PART D — How Long Data Is Kept')
d.table(
    ['Data Type', 'How Long Kept', 'Who Can Delete It'],
    [
        ['Live skeleton/movement data', '90 days (rolling)', 'Account holder at any time'],
        ['Fall events and alerts', 'Duration of monitoring + 2 years', 'Account holder — contact EchoPose to delete earlier'],
        ['Physiological estimates', '90 days (rolling)', 'Account holder at any time'],
        ['Activity summaries', '90 days (rolling)', 'Account holder at any time'],
    ]
)

d.h1('PART E — Your Rights')
d.para('You have the following rights at any time:')
d.bullet([
    'Right to withdraw consent — you can stop monitoring at any time by asking the account holder or contacting ' + EMAIL,
    'Right to access — you can request a copy of all data held about you',
    'Right to erasure — you can ask for your data to be deleted',
    'Right to object — you can object to specific types of processing',
    'Right to complain to the ICO — ico.org.uk — Tel: 0303 123 1113',
])

d.h1('PART F — Consent Declaration')
d.para(
    'I confirm that I have read and understood the information above. '
    'I understand what EchoPose monitors, who will see my data, and how long it will be kept. '
    'I understand that EchoPose is NOT a medical device and is NOT a substitute for emergency services. '
    'In an emergency, I understand that 999 must be called.'
)
d.para('')

d.banner('Section F1 — Consent by the Monitored Person (if they have capacity)', TEAL)
d.para('I, the person who will be monitored, give my free and informed consent to be monitored by EchoPose as described above.')
d.para('')
d.para('Full Name:  _______________________________________________________________')
d.para('')
d.para('Date of Birth:  ___________________________________________________________')
d.para('')
d.para('Address:  ________________________________________________________________')
d.para('')
d.para('Signature:  _______________________________________________________________')
d.para('')
d.para('Date:  ___________________________________________________________________')
d.para('')

d.banner('Section F2 — Consent by Legal Representative (if person lacks capacity)', ORANGE)
d.callout(
    'Complete this section ONLY if the monitored person lacks mental capacity to consent for themselves '
    '(e.g. advanced dementia). The representative must have legal authority (Lasting Power of Attorney for Health & Welfare '
    'or deputyship order). A copy of the authority document must be kept on file.',
    ORANGE, '⚠ NOTE'
)
d.para('')
d.para('I confirm that I have legal authority to consent on behalf of the monitored person.')
d.para('')
d.para('Representative Full Name:  ________________________________________________')
d.para('')
d.para('Relationship to Monitored Person:  ________________________________________')
d.para('')
d.para('Legal Authority (e.g. LPA reference):  _____________________________________')
d.para('')
d.para('Signature:  _______________________________________________________________')
d.para('')
d.para('Date:  ___________________________________________________________________')
d.para('')

d.banner('Section F3 — Account Holder / Carer Declaration', NAVY)
d.para(f'I confirm that I have explained the EchoPose system to the monitored person (or their representative), '
       f'and that this consent form has been completed voluntarily.')
d.para('')
d.para('Account Holder Full Name:  ________________________________________________')
d.para('')
d.para('Organisation (if applicable):  ____________________________________________')
d.para('')
d.para('Signature:  _______________________________________________________________')
d.para('')
d.para('Date:  ___________________________________________________________________')
d.para('')
d.para(f'Return completed form to:  {EMAIL}  or keep a physical copy on file.')

d.footer(f'EchoPose Consent Form · {OWNER} · {DATE} · UK GDPR Article 9')
d.save()


# ══════════════════════════════════════════════════════════════════════════════
#  5.  RESELLER / CHANNEL PARTNER AGREEMENT
# ══════════════════════════════════════════════════════════════════════════════
print('Building Reseller / Channel Partner Agreement...')
d = Doc('EchoPose_Reseller_Channel_Partner_Agreement.docx', BLUE)
d.cover(
    'Reseller & Channel Partner Agreement',
    'For smart home installers, care tech consultants, and integration partners',
    'COMMERCIAL LEGAL TEMPLATE  |  Version 1.0'
)

d.callout(
    'Use this agreement when appointing a business or individual to sell or install '
    'EchoPose on your behalf in exchange for a commission or margin.',
    BLUE, 'WHEN TO USE THIS DOCUMENT'
)

d.para('RESELLER AND CHANNEL PARTNER AGREEMENT', bold=True, size=14, color=NAVY)
d.para('')
d.para('This Agreement is entered into between:', size=11)
d.para('')
d.para(f'Principal:  {OWNER}, trading as {TRADING_AS}', bold=True)
d.para(f'            Contact: {EMAIL}')
d.para('')
d.para('Reseller:   ___________________________________________________________', bold=True)
d.para('            Company / Trading Name:  _______________________________________')
d.para('            Registered Address:  ____________________________________________')
d.para('            Contact Name:  __________________________________________________')
d.para('            Email:  ________________________________________________________')
d.para('')
d.para('Together referred to as "the Parties".')

d.h1('1.  Appointment')
d.para(
    f'EchoPose appoints the Reseller as a non-exclusive authorised reseller of EchoPose products '
    f'and services within the following territory: __________________________________ ("the Territory"). '
    f'The Reseller accepts this appointment on the terms set out in this Agreement. '
    f'"Non-exclusive" means EchoPose may appoint additional resellers in the same territory without restriction.'
)

d.h1('2.  Products and Services Covered')
d.table(
    ['Product / Service', 'Reseller May Sell?', 'Notes'],
    [
        ['EchoPose Home subscription (£49/mo)', 'Yes', 'Reseller bills customer and pays EchoPose net price'],
        ['EchoPose Care subscription (£199/mo)', 'Yes', ''],
        ['EchoPose Pro subscription (£499/mo)', 'Yes', ''],
        ['Hardware Starter Kit', 'Yes', 'Reseller purchases from EchoPose at trade price'],
        ['Professional Installation Service', 'Yes — Reseller performs installation', 'Must complete EchoPose installer certification'],
        ['On-Premise License', 'Yes — with prior written approval per deal', 'Each deal requires EchoPose sign-off'],
        ['EchoPose branding / white-labelling', 'No', 'Not permitted without separate written agreement'],
    ]
)

d.pb()
d.h1('3.  Pricing, Margin and Commission')
d.h2('3.1  Reseller Pricing Model')
d.para(
    'The Reseller may choose one of the following commercial models, '
    'as agreed in writing with EchoPose at the time of appointment:'
)
d.table(
    ['Model', 'How It Works', 'Reseller Margin / Commission'],
    [
        ['Referral Model', 'Reseller refers customers to EchoPose. EchoPose bills the customer directly.',
         '15% of first 12 months\' subscription revenue per referred customer'],
        ['Reseller Model', 'Reseller bills the customer. Reseller pays EchoPose a net price.',
         'Reseller sets their own price. EchoPose net price = 75% of EchoPose list price.'],
        ['Installation-Only', 'Reseller provides installation service only. Customer subscribes directly.',
         '£100–200 per installation (agreed per job)'],
    ]
)

d.h2('3.2  Payment Terms')
d.para(
    'Under the Referral Model: EchoPose pays commission monthly in arrears, '
    'within 30 days of receiving cleared customer payment. '
    'Under the Reseller Model: Reseller pays EchoPose net price within 30 days of invoice. '
    'Late payment accrues interest at 8% above Bank of England base rate per annum '
    'under the Late Payment of Commercial Debts Act 1998.'
)

d.h2('3.3  Price Changes')
d.para(
    'EchoPose may adjust list prices or net prices with 60 days written notice to the Reseller. '
    'Existing customer subscriptions are honoured at the original price until renewal or upgrade.'
)

d.h1('4.  Reseller Obligations')
d.para('The Reseller agrees to:')
d.numbered([
    'Actively promote EchoPose products and services within the Territory and represent them accurately;',
    'Not make any representations about EchoPose that are false, misleading, or inconsistent with EchoPose\'s official marketing materials;',
    'NOT describe EchoPose as a medical device, clinical monitoring solution, or substitute for emergency services;',
    'Ensure that all customers are provided with and agree to EchoPose\'s Terms of Service and Privacy Policy before activation;',
    'Ensure that a completed EchoPose Consent Form is obtained for every monitored individual before the system is activated;',
    'Complete EchoPose installer certification before performing any hardware installation (certification provided free of charge online);',
    'Provide first-line customer support for their own customers; escalate unresolved technical issues to EchoPose;',
    'Not engage in any activity that brings EchoPose\'s reputation into disrepute;',
    'Maintain adequate professional indemnity and public liability insurance (minimum £1 million cover);',
    'Not appoint sub-resellers without EchoPose\'s prior written consent.',
])

d.h1('5.  EchoPose Obligations')
d.para('EchoPose agrees to:')
d.numbered([
    'Provide the Reseller with up-to-date marketing materials, product documentation, and pricing;',
    'Provide free online installer certification training;',
    'Provide second-line technical support to the Reseller (not directly to Reseller\'s customers) within 2 business days;',
    'Pay commission or honour agreed net pricing as set out in Clause 3;',
    'Give reasonable advance notice of product changes, discontinuations, or price changes.',
])

d.pb()
d.h1('6.  Branding and Marketing')
d.para(
    'EchoPose grants the Reseller a limited, non-exclusive licence to use the EchoPose name and logo '
    'solely for the purpose of promoting EchoPose products under this Agreement. '
    'The Reseller must not alter the logo or create derivative marks. '
    'Any marketing materials using EchoPose branding must be approved in writing by EchoPose before use. '
    'The Reseller may not imply any endorsement of the Reseller\'s own brand by EchoPose without written consent.'
)

d.h1('7.  Intellectual Property')
d.para(
    'All intellectual property in EchoPose products, software, firmware, algorithms, and documentation '
    f'remains the exclusive property of {OWNER}. '
    'This Agreement does not transfer any IP rights to the Reseller. '
    'The Reseller shall immediately notify EchoPose of any suspected infringement of EchoPose IP '
    'that comes to their attention.'
)

d.h1('8.  Confidentiality')
d.para(
    'Each Party agrees to keep confidential all proprietary and commercially sensitive information '
    'disclosed by the other Party in connection with this Agreement, '
    'including but not limited to pricing, customer lists, technology, and business plans. '
    'This obligation continues for 3 years after termination of this Agreement. '
    'A separate NDA may be required for disclosure of technical details of EchoPose\'s AI engine or signal processing.'
)

d.h1('9.  Data Protection')
d.para(
    'The Reseller must comply with all applicable data protection laws in connection with its activities under this Agreement. '
    'Where the Reseller handles personal data of EchoPose customers or monitored individuals, '
    'it does so as an independent data controller and is solely responsible for its own compliance. '
    'The Reseller must obtain appropriate consent forms from all monitored individuals '
    'and must not activate any EchoPose system without a completed consent form. '
    'The Reseller must enter into a DPA with EchoPose if it handles personal data on EchoPose\'s behalf.'
)

d.h1('10.  Liability')
d.para(
    'The Reseller is solely liable for any claims arising from: '
    'misrepresentation of EchoPose products (including describing them as medical devices); '
    'improper installation; failure to obtain required consent from monitored individuals; '
    'breach of applicable law in their territory. '
    'EchoPose\'s liability to the Reseller is limited to the net amounts paid by the Reseller '
    'to EchoPose in the 3 months preceding the claim. '
    'EchoPose is not liable for the Reseller\'s loss of profit or indirect losses.'
)

d.h1('11.  Term and Termination')
d.para(
    'This Agreement commences on the Effective Date and continues for 12 months, '
    'automatically renewing for successive 12-month periods unless terminated. '
    'Either Party may terminate this Agreement by giving 60 days written notice at any time. '
    'EchoPose may terminate immediately for: material breach by Reseller; insolvency of Reseller; '
    'or any act by Reseller that brings EchoPose into disrepute. '
    'On termination, the Reseller must immediately cease using EchoPose branding and marketing materials, '
    'and must transfer all customer contacts to EchoPose within 14 days.'
)

d.h1('12.  Governing Law')
d.para(
    f'This Agreement is governed by the laws of {COUNTRY}. '
    'Any disputes shall be subject to the exclusive jurisdiction of the courts of England and Wales.'
)

d.h1('13.  Entire Agreement')
d.para(
    'This Agreement constitutes the entire agreement between the Parties with respect to its subject matter '
    'and supersedes all prior agreements, negotiations, and understandings. '
    'Any variation must be agreed in writing and signed by both Parties.'
)

d.h1('14.  Signatures')
d.para('')
d.para('ECHOPOSE (Principal):', bold=True)
d.sig(OWNER, f'Owner, {TRADING_AS}', witness=True)
d.divider()
d.para('RESELLER:', bold=True)
d.sig('___________________________________________', 'Authorised Signatory', company='Reseller', witness=True)

d.callout(
    'This Agreement should be reviewed by a qualified commercial solicitor before use. '
    'It is drafted under English law as a starting framework only.',
    ORANGE, '⚠ LEGAL NOTE'
)
d.footer(f'EchoPose Reseller Agreement · {OWNER} · {DATE} · Governing Law: {COUNTRY}')
d.save()

print(f'\nAll 5 legal/operational documents saved to {OUT}/')
